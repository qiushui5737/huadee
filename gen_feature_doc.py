"""
生成政民互动板块功能开发与实现文档（docx格式）
输出路径：E:\huadee\政民互动功能开发与实现.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('政府网站集约化平台')
run.font.name = '黑体'
run.font.size = Pt(28)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('政民互动板块')
run.font.name = '黑体'
run.font.size = Pt(24)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('功能开发与实现文档')
run.font.name = '黑体'
run.font.size = Pt(22)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    '开发周期：2026年7月12日 — 2026年7月20日',
    '技术栈：SpringBoot 3.2.5 + Vue3 + MyBatis-Plus + MySQL 8.0',
    '功能模块：咨询、建议、投诉、省长信箱、意见征集',
    '接口数量：46个RESTful API',
    '数据表：9张业务表',
    '前端页面：C端13个 + B端5个',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = '宋体'
    run.font.size = Pt(13)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
add_title('目  录', level=1)
toc_items = [
    '一、概述',
    '二、我要咨询功能开发与实现',
    '    2.1 功能描述',
    '    2.2 业务流程',
    '    2.3 核心代码',
    '    2.4 前端实现',
    '三、我要建议功能开发与实现',
    '    3.1 功能描述',
    '    3.2 业务流程',
    '    3.3 核心代码',
    '    3.4 前端实现',
    '四、我要投诉功能开发与实现',
    '    4.1 功能描述',
    '    4.2 业务流程',
    '    4.3 核心代码',
    '    4.4 前端实现',
    '五、省长信箱功能开发与实现',
    '    5.1 功能描述',
    '    5.2 业务流程',
    '    5.3 核心代码',
    '    5.4 前端实现',
    '六、意见征集功能开发与实现',
    '    6.1 功能描述',
    '    6.2 业务流程',
    '    6.3 核心代码',
    '    6.4 前端实现',
    '七、接口汇总',
    '八、实现截图说明',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 一、概述
# ============================================================
add_title('一、概述', level=1)

add_body('政民互动板块是政府网站集约化平台的核心模块之一，旨在搭建政府与群众之间的沟通桥梁。本板块包含五个功能子系统：我要咨询、我要建议、我要投诉、省长信箱和意见征集。')

add_body('五个功能子系统各自面向不同的政民互动场景：')

add_bullet('我要咨询：群众就政策法规、办事流程等问题向政府部门咨询，管理员受理答复，系统生成ZX开头咨询单号，答复期限20个自然日。')
add_bullet('我要建议：群众对政府网站、部门服务、政策制定等提出意见建议，支持4种建议类型（网站建议/部门建议/我要纠错/助企惠企服务专区建议），系统生成JY开头建议单号。')
add_bullet('我要投诉：群众对政府部门的服务、行政行为等进行投诉举报，支持三级投诉单位和三级问题分类，系统生成TS开头投诉单号。')
add_bullet('省长信箱：群众通过省长信箱向省级政府反映问题，采用邮件式界面写信，支持多轮对话交流、分派管理、督办管理、满意度评价，系统生成LX开头信件单号。')
add_bullet('意见征集：政府发布征集活动，群众提交意见，管理员回复意见（采纳/部分采纳/不采纳），最终公示结果反馈，形成完整的"发布→收集→处理→反馈"闭环。')

add_body('技术架构方面，后端采用SpringBoot 3.2.5 + MyBatis-Plus 3.5.5 + MySQL 8.0 + Redis + JWT认证，前端采用Vue3 + TypeScript + Element Plus + Vite 5，通过RESTful API实现前后端数据交互。')

add_body('本板块共开发8个Controller、8个Service、15个Entity、15个Mapper，46个RESTful接口，C端13个页面、B端5个管理页面，9张业务数据表。')

doc.add_page_break()

# ============================================================
# 二、我要咨询
# ============================================================
add_title('二、我要咨询功能开发与实现', level=1)

add_title('2.1 功能描述', level=2)

add_body('我要咨询功能为群众提供政策法规、办事流程等问题的咨询渠道。群众在前端填写咨询表单提交后，系统自动生成ZX开头的咨询单号，管理员在后台查看咨询列表、进行答复、修改状态、办结咨询。群众可通过咨询单号查询处理进度和答复内容，仅本人可查询自己的咨询记录。')

add_body('该功能包含以下7个功能点：')

add_bullet('提交咨询（C端）：群众填写真实姓名、手机号码、咨询主题、咨询市州、咨询内容等必填项，可选填电话号码、电子邮箱、事件发生地（区/街道/地址）、附件图片，选择是否公开或保密。系统校验必填字段后生成ZX单号，设置状态为"待受理"，计算20天答复期限，保存到consultation表。')
add_bullet('查询进度（C端）：群众输入咨询单号查询处理状态和答复内容。后端同时匹配单号和当前用户ID进行权限校验，非本人查询返回403错误。')
add_bullet('咨询列表（B端）：管理员分页查看咨询列表，支持按关键词（姓名/单号/主题模糊搜索）、状态（待受理/处理中/已答复/已办结）、城市进行筛选，按创建时间倒序排列。')
add_bullet('答复咨询（B端）：管理员填写答复内容和答复人姓名，提交后状态自动变为"已答复"，记录答复时间和答复人。')
add_bullet('修改状态（B端）：管理员手动修改咨询状态，后端使用白名单校验状态值必须为"待受理/处理中/已答复/已办结"之一，防止非法状态值写入。')
add_bullet('办结咨询（B端）：管理员将已答复的咨询标记为已办结，前置校验当前状态必须为"已答复"，防止跳过答复直接办结。')
add_bullet('数据统计（B端）：管理员查看各状态咨询数量和超时未答复数量，用于工作监控和绩效考核。')

add_title('2.2 业务流程', level=2)

add_body('咨询功能的完整业务流程如下：')

add_bullet('群众登录系统后进入咨询提交页面，填写表单信息。')
add_bullet('前端进行基础校验（必填项检查），通过后调用POST /api/v1/consultation接口提交。')
add_bullet('后端ConsultationController接收请求，从JWT Token中提取userId，解析表单字段构建Consultation实体。')
add_bullet('ConsultationService.submit()方法执行参数校验、生成ZX单号、设置初始状态和答复期限、保存到数据库。')
add_bullet('管理员在后台管理页面查看咨询列表，使用关键词/状态/城市筛选目标咨询。')
add_bullet('管理员点击答复按钮，填写答复内容，调用POST /api/v1/consultation/{consultNo}/reply接口。')
add_bullet('后端更新咨询状态为"已答复"，记录答复内容和时间。')
add_bullet('管理员可通过修改状态接口手动调整状态流转，或调用办结接口完成闭环。')
add_bullet('群众通过单号查询接口查看处理进度和答复结果。')

add_title('2.3 核心代码', level=2)

add_body('以下为咨询功能后端Service层的核心实现代码：', bold=True)

add_body('提交咨询方法：', bold=True)
add_code('public Consultation submit(Consultation c) {')
add_code('    // 参数校验')
add_code('    if (!StringUtils.hasText(c.getRealName()))')
add_code('        throw BusinessException.of("真实姓名不能为空");')
add_code('    if (!StringUtils.hasText(c.getPhone()))')
add_code('        throw BusinessException.of("手机号码不能为空");')
add_code('    if (!StringUtils.hasText(c.getTitle()))')
add_code('        throw BusinessException.of("咨询主题不能为空");')
add_code('    if (!StringUtils.hasText(c.getCity()))')
add_code('        throw BusinessException.of("请选择咨询市州");')
add_code('    if (!StringUtils.hasText(c.getContent()))')
add_code('        throw BusinessException.of("咨询内容不能为空");')
add_code('    // 生成咨询单号: ZX + 时间戳 + 4位随机数')
add_code('    String no = "ZX" + LocalDateTime.now()')
add_code('        .format(DateTimeFormatter.ofPattern("yyyyMMddHHmmss"))')
add_code('        + String.format("%04d", ThreadLocalRandom.current().nextInt(10000));')
add_code('    c.setConsultNo(no);')
add_code('    c.setStatus("待受理");')
add_code('    c.setDeadline(LocalDate.now().plusDays(20));')
add_code('    save(c);')
add_code('    return c;')
add_code('}')

add_body('进度查询方法（仅本人权限）：', bold=True)
add_code('public Consultation getByConsultNoAndUserId(String consultNo, Long userId) {')
add_code('    LambdaQueryWrapper<Consultation> wrapper = new LambdaQueryWrapper<>();')
add_code('    wrapper.eq(Consultation::getConsultNo, consultNo)')
add_code('           .eq(Consultation::getUserId, userId);')
add_code('    Consultation c = getOne(wrapper);')
add_code('    if (c == null) {')
add_code('        throw BusinessException.of(403, "咨询单号不存在或无权查看");')
add_code('    }')
add_code('    return c;')
add_code('}')

add_body('答复咨询方法：', bold=True)
add_code('public void reply(String consultNo, String replyContent, String operator) {')
add_code('    Consultation c = getByConsultNo(consultNo);')
add_code('    if (!StringUtils.hasText(replyContent))')
add_code('        throw BusinessException.of("答复内容不能为空");')
add_code('    c.setStatus("已答复");')
add_code('    c.setReplyContent(replyContent);')
add_code('    c.setReplyBy(operator);')
add_code('    c.setReplyTime(LocalDateTime.now());')
add_code('    updateById(c);')
add_code('}')

add_body('办结咨询方法（前置状态校验）：', bold=True)
add_code('public void finish(String consultNo) {')
add_code('    Consultation c = getByConsultNo(consultNo);')
add_code('    if (!"已答复".equals(c.getStatus()))')
add_code('        throw BusinessException.of("只有已答复状态的咨询才能办结");')
add_code('    c.setStatus("已办结");')
add_code('    updateById(c);')
add_code('}')

add_title('2.4 前端实现', level=2)

add_body('前端API层封装了7个咨询相关接口函数：submitConsultation（提交）、consultationProgress（进度查询）、consultationList（列表）、replyConsultation（答复）、updateConsultationStatus（改状态）、finishConsultation（办结）、consultationStats（统计）。')

add_body('C端页面包括Consultation.vue（咨询提交表单，包含姓名/电话/主题/市州/内容等字段，支持附件上传和公开/保密选项）和ConsultationDetail.vue（咨询详情与答复展示）。B端页面为ConsultationMgmt.vue（咨询管理列表，支持筛选、答复弹窗、状态修改、办结操作、统计面板）。')

doc.add_page_break()

# ============================================================
# 三、我要建议
# ============================================================
add_title('三、我要建议功能开发与实现', level=1)

add_title('3.1 功能描述', level=2)

add_body('我要建议功能为群众提供建言献策渠道，支持对政府网站、部门服务、政策制定等提出意见建议。与咨询功能相比，建议功能增加了建议类型分类（网站建议/部门建议/我要纠错/助企惠企服务专区建议）和身份证号字段，其余业务流程与咨询功能保持一致。')

add_body('该功能包含以下7个功能点：')

add_bullet('提交建议（C端）：群众填写建议标题、选择建议类型（4种）、姓名、身份证号、电子邮箱、联系电话、联系地址、省市区地址、建议内容、附件上传，选择是否公开或保密。系统校验后生成JY单号，设置状态为"待受理"，计算20天答复期限。')
add_bullet('查询进度（C端）：通过建议单号查询处理状态和答复内容，仅本人可查。')
add_bullet('建议列表（B端）：管理员分页查看建议列表，支持关键词/状态/城市筛选。')
add_bullet('答复建议（B端）：管理员填写答复内容，状态自动变为"已答复"。')
add_bullet('修改状态（B端）：手动修改建议状态，白名单校验状态值合法性。')
add_bullet('办结建议（B端）：将已答复的建议标记为已办结，前置校验状态。')
add_bullet('数据统计（B端）：查看各状态数量和超时数量。')

add_title('3.2 业务流程', level=2)

add_body('建议功能的业务流程与咨询功能高度一致，主要差异在于：')

add_bullet('表单字段更多：增加了建议类型选择（4种类型下拉框）、身份证号、省市区三级地址选择。')
add_bullet('单号前缀不同：建议单号以JY开头，咨询单号以ZX开头。')
add_bullet('数据表不同：建议数据存储在suggestion表，咨询数据存储在consultation表。')
add_bullet('其余流程（提交→受理→答复→办结→查询）完全相同，采用统一的Controller-Service设计模式。')

add_title('3.3 核心代码', level=2)

add_body('提交建议方法：', bold=True)
add_code('public Suggestion submit(Suggestion s) {')
add_code('    if (!StringUtils.hasText(s.getTitle()))')
add_code('        throw BusinessException.of("建议标题不能为空");')
add_code('    if (!StringUtils.hasText(s.getRealName()))')
add_code('        throw BusinessException.of("建议人姓名不能为空");')
add_code('    if (!StringUtils.hasText(s.getPhone()))')
add_code('        throw BusinessException.of("联系电话不能为空");')
add_code('    if (!StringUtils.hasText(s.getContent()))')
add_code('        throw BusinessException.of("建议内容不能为空");')
add_code('    // 生成建议单号: JY + 时间戳 + 4位随机数')
add_code('    String no = "JY" + LocalDateTime.now()')
add_code('        .format(DateTimeFormatter.ofPattern("yyyyMMddHHmmss"))')
add_code('        + String.format("%04d", ThreadLocalRandom.current().nextInt(10000));')
add_code('    s.setSuggestNo(no);')
add_code('    s.setStatus("待受理");')
add_code('    s.setDeadline(LocalDate.now().plusDays(20));')
add_code('    save(s);')
add_code('    return s;')
add_code('}')

add_body('答复建议方法：', bold=True)
add_code('public void reply(String suggestNo, String replyContent, String operator) {')
add_code('    Suggestion s = getBySuggestNo(suggestNo);')
add_code('    if (!StringUtils.hasText(replyContent))')
add_code('        throw BusinessException.of("答复内容不能为空");')
add_code('    s.setStatus("已答复");')
add_code('    s.setReplyContent(replyContent);')
add_code('    s.setReplyBy(operator);')
add_code('    s.setReplyTime(LocalDateTime.now());')
add_code('    updateById(s);')
add_code('}')

add_title('3.4 前端实现', level=2)

add_body('前端API层封装了7个建议相关接口函数：submitSuggestion、suggestionProgress、suggestionList、replySuggestion、updateSuggestionStatus、finishSuggestion、suggestionStats。')

add_body('C端页面包括SubmitSuggest.vue（建议提交表单，包含建议类型下拉选择、三级地址选择器）和SuggestionDetail.vue（建议详情展示）。B端页面为SuggestionMgmt.vue（建议管理列表，功能布局与咨询管理一致）。')

doc.add_page_break()

# ============================================================
# 四、我要投诉
# ============================================================
add_title('四、我要投诉功能开发与实现', level=1)

add_title('4.1 功能描述', level=2)

add_body('我要投诉功能为群众提供投诉监督渠道，支持对政府部门的服务、行政行为等进行投诉举报。该功能是三个统一设计模式功能（咨询/建议/投诉）中字段最丰富的，包含三级投诉单位层级和三级问题类型层级，共计30余个字段。')

add_body('该功能包含以下7个功能点：')

add_bullet('提交投诉（C端）：群众填写投诉标题、选择三级投诉单位（市州→部门→子部门联动选择）、选择三级问题类型、填写投诉人信息（姓名/身份证/邮箱/电话/地址）、事件发生地（省/市/区/具体地址）、投诉内容、附件上传，选择是否公开或保密。系统校验必填字段后生成TS单号，设置状态为"待受理"，计算20天答复期限。')
add_bullet('查询进度（C端）：通过投诉单号查询处理状态和答复内容，仅本人可查。')
add_bullet('投诉列表（B端）：管理员分页查看投诉列表，支持关键词/状态/城市筛选。')
add_bullet('答复投诉（B端）：管理员填写答复内容，状态自动变为"已答复"。')
add_bullet('修改状态（B端）：手动修改投诉状态，白名单校验状态值。')
add_bullet('办结投诉（B端）：将已答复的投诉标记为已办结，前置校验状态。')
add_bullet('数据统计（B端）：查看各状态数量和超时数量。')

add_title('4.2 业务流程', level=2)

add_body('投诉功能的业务流程与咨询、建议功能保持一致的统一模式，主要差异在于：')

add_bullet('表单字段最丰富：包含三级投诉单位联动选择（选择市州后加载对应部门，选择部门后加载子部门）和三级问题类型联动选择，以及完整的事件发生地信息。')
add_bullet('单号前缀为TS，数据存储在complaint表。')
add_bullet('投诉单位字段（complaint_unit_l1/l2/l3）和问题类型字段（problem_type_l1/l2/l3）采用三级层级结构，便于后续按层级统计分析。')

add_title('4.3 核心代码', level=2)

add_body('提交投诉方法：', bold=True)
add_code('public Complaint submit(Complaint c) {')
add_code('    if (!StringUtils.hasText(c.getTitle()))')
add_code('        throw BusinessException.of("投诉标题不能为空");')
add_code('    if (!StringUtils.hasText(c.getRealName()))')
add_code('        throw BusinessException.of("投诉人姓名不能为空");')
add_code('    if (!StringUtils.hasText(c.getPhone()))')
add_code('        throw BusinessException.of("联系电话不能为空");')
add_code('    if (!StringUtils.hasText(c.getContent()))')
add_code('        throw BusinessException.of("投诉内容不能为空");')
add_code('    if (!StringUtils.hasText(c.getComplaintUnitL1()))')
add_code('        throw BusinessException.of("请选择投诉单位");')
add_code('    // 生成投诉单号: TS + 时间戳 + 4位随机数')
add_code('    String no = "TS" + LocalDateTime.now()')
add_code('        .format(DateTimeFormatter.ofPattern("yyyyMMddHHmmss"))')
add_code('        + String.format("%04d", ThreadLocalRandom.current().nextInt(10000));')
add_code('    c.setComplaintNo(no);')
add_code('    c.setStatus("待受理");')
add_code('    c.setDeadline(LocalDate.now().plusDays(20));')
add_code('    save(c);')
add_code('    return c;')
add_code('}')

add_body('状态修改方法（白名单校验）：', bold=True)
add_code('public void updateStatus(String complaintNo, String status) {')
add_code('    Complaint c = getByComplaintNo(complaintNo);')
add_code('    String[] validStatuses = {"待受理", "处理中", "已答复", "已办结"};')
add_code('    boolean valid = false;')
add_code('    for (String st : validStatuses) {')
add_code('        if (st.equals(status)) { valid = true; break; }')
add_code('    }')
add_code('    if (!valid)')
add_code('        throw BusinessException.of("无效的状态值: " + status);')
add_code('    c.setStatus(status);')
add_code('    updateById(c);')
add_code('}')

add_title('4.4 前端实现', level=2)

add_body('前端API层封装了7个投诉相关接口函数：submitComplaint、complaintProgress、complaintList、replyComplaint、updateComplaintStatus、finishComplaint、complaintStats。')

add_body('C端页面为SubmitComplaint.vue（投诉提交表单，包含三级投诉单位联动选择器、三级问题类型联动选择器、完整的事件发生地填写区域，是三个功能中表单最复杂的页面）。B端页面为ComplaintMgmt.vue（投诉管理列表）。')

doc.add_page_break()

# ============================================================
# 五、省长信箱
# ============================================================
add_title('五、省长信箱功能开发与实现', level=1)

add_title('5.1 功能描述', level=2)

add_body('省长信箱是政民互动板块中功能最复杂、业务场景最丰富的子系统。与咨询/建议/投诉的简单问答模式不同，省长信箱采用邮件式界面和多轮对话模式，支持群众与管理员之间的多次交流，同时提供分派管理、督办管理、满意度评价等高级功能。')

add_body('该功能涉及4张数据表：message（信件主表）、message_reply（多轮对话记录表）、message_assignment（分派记录表）、message_supervision（督办记录表）。')

add_body('该功能包含以下11个功能点：')

add_bullet('提交信件（C端）：群众采用邮件式界面写信，填写收件人（目标部门）、主题、富文本正文，选择是否公开。系统生成LX开头信件单号，设置状态为"待分派"，默认处理期限7天。')
add_bullet('信件列表（C端/B端）：C端显示公开信件和本人信件（私密信件仅本人可见），B端显示全部信件。支持按关键词、状态（待分派/已分派/已回复/已办结）、类型（咨询/投诉/建议/求助）筛选。')
add_bullet('信件详情（C端/B端）：展示信件完整信息，包含多轮对话记录。对话记录按时间升序排列，区分管理员（ADMIN）和群众（USER）身份，使用不同的头像和消息气泡样式展示。')
add_bullet('管理员回复（B端）：管理员填写回复内容，回复记录写入message_reply表（user_type=ADMIN），同时更新主表的最新回复信息（保持向后兼容），状态变为"已回复"。')
add_bullet('群众追问（C端）：仅信件本人可追问，后端校验userId与信件userId一致，回复写入message_reply表（user_type=USER）。非本人访问时隐藏回复输入区域，仅可查看对话记录。')
add_bullet('分派管理（B端）：管理员将信件分派到目标部门，更新dept_code字段，状态变为"已分派"。')
add_bullet('督办管理（B端）：管理员设置督办标记和督办时间，用于超时预警和重点跟踪。')
add_bullet('满意度评价（C端）：群众对管理员回复进行1-5星评价，可填写评价内容，评价数据存储在message表的rating和rating_content字段。')
add_bullet('办结信件（B端）：管理员标记信件为已办结。')
add_bullet('热门信件（C端）：展示最近10条已回复或已办结的公开信件，用于首页推荐展示。')
add_bullet('数据统计（B端）：查看各状态信件数量和超时数量。')

add_title('5.2 业务流程', level=2)

add_body('省长信箱的完整业务流程如下：')

add_bullet('群众登录系统后进入写信页面（/interaction/write），页面采用邮件客户端风格设计，包含收件人选择、主题输入、富文本工具栏、HTML正文编辑区。')
add_bullet('前端校验标题和内容必填后，调用POST /api/v1/interaction/message接口提交。')
add_bullet('后端MessageService.submit()方法生成LX单号，设置状态为"待分派"，保存信件。')
add_bullet('管理员在后台查看信件列表，可执行分派操作（选择目标部门）、回复操作（填写回复内容）、督办操作（设置督办期限）。')
add_bullet('管理员回复时，后端同时执行两个操作：将回复记录插入message_reply表（多轮对话），更新message主表的最新回复信息（向后兼容）。')
add_bullet('群众查看信件详情时，后端查询message表获取信件信息，同时查询message_reply表获取全部对话记录，按时间升序返回。')
add_bullet('群众追问时，后端校验当前用户是否为信件本人（userId匹配），校验通过后插入message_reply表（user_type=USER）。')
add_bullet('群众可对已回复的信件进行满意度评价（1-5星）。')
add_bullet('管理员办结信件，流程结束。')

add_title('5.3 核心代码', level=2)

add_body('提交信件方法：', bold=True)
add_code('public Message submit(Message message) {')
add_code('    if (!StringUtils.hasText(message.getTitle()))')
add_code('        throw BusinessException.of("留言标题不能为空");')
add_code('    if (!StringUtils.hasText(message.getContent()))')
add_code('        throw BusinessException.of("留言内容不能为空");')
add_code('    if (!StringUtils.hasText(message.getType()))')
add_code('        message.setType("咨询");')
add_code('    message.setStatus("待分派");')
add_code('    if (message.getIsPublic() == null)')
add_code('        message.setIsPublic(true);')
add_code('    // 生成信件单号: LX + 时间戳 + 4位随机数')
add_code('    String no = "LX" + LocalDateTime.now()')
add_code('        .format(DateTimeFormatter.ofPattern("yyyyMMddHHmmss"))')
add_code('        + String.format("%04d", ThreadLocalRandom.current().nextInt(10000));')
add_code('    message.setConsultNo(no);')
add_code('    message.setDeadline(LocalDateTime.now().plusDays(7));')
add_code('    save(message);')
add_code('    return message;')
add_code('}')

add_body('管理员回复方法（多轮对话）：', bold=True)
add_code('public void reply(Long id, String content, String operator) {')
add_code('    Message message = getById(id);')
add_code('    if (message == null)')
add_code('        throw BusinessException.of("留言不存在");')
add_code('    if (!StringUtils.hasText(content))')
add_code('        throw BusinessException.of("回复内容不能为空");')
add_code('    // 插入回复记录到对话表')
add_code('    MessageReply reply = new MessageReply();')
add_code('    reply.setMessageId(id);')
add_code('    reply.setUserName(operator);')
add_code('    reply.setUserType("ADMIN");')
add_code('    reply.setContent(content);')
add_code('    reply.setCreateTime(LocalDateTime.now());')
add_code('    messageReplyMapper.insert(reply);')
add_code('    // 同时更新主表的最新回复信息')
add_code('    message.setReplyContent(content);')
add_code('    message.setReplyBy(operator);')
add_code('    message.setReplyTime(LocalDateTime.now());')
add_code('    message.setStatus("已回复");')
add_code('    updateById(message);')
add_code('}')

add_body('群众追问方法（本人权限校验）：', bold=True)
add_code('public void userReply(Long id, Long userId, String userName, String content) {')
add_code('    Message message = getById(id);')
add_code('    if (message == null)')
add_code('        throw BusinessException.of("留言不存在");')
add_code('    // 校验是否为信件本人')
add_code('    if (userId == null || !userId.equals(message.getUserId()))')
add_code('        throw BusinessException.of(403, "只有信件本人才能回复");')
add_code('    if (!StringUtils.hasText(content))')
add_code('        throw BusinessException.of("回复内容不能为空");')
add_code('    MessageReply reply = new MessageReply();')
add_code('    reply.setMessageId(id);')
add_code('    reply.setUserId(userId);')
add_code('    reply.setUserName(userName);')
add_code('    reply.setUserType("USER");')
add_code('    reply.setContent(content);')
add_code('    reply.setCreateTime(LocalDateTime.now());')
add_code('    messageReplyMapper.insert(reply);')
add_code('}')

add_body('私密信件可见性校验：', bold=True)
add_code('public Message detail(Long id, Long currentUserId) {')
add_code('    Message message = getById(id);')
add_code('    if (message == null)')
add_code('        throw BusinessException.of("留言不存在");')
add_code('    // 私密信件只有本人和管理员可查看')
add_code('    if (Boolean.FALSE.equals(message.getIsPublic())) {')
add_code('        if (currentUserId == null || !currentUserId.equals(message.getUserId()))')
add_code('            throw BusinessException.of(403, "该信件为私密信件，无权查看");')
add_code('    }')
add_code('    return message;')
add_code('}')

add_body('C端列表可见性过滤（公开+本人）：', bold=True)
add_code('public IPage<Message> listPage(..., Long currentUserId, Boolean isPublicOnly, ...) {')
add_code('    LambdaQueryWrapper<Message> wrapper = new LambdaQueryWrapper<>();')
add_code('    if (Boolean.TRUE.equals(isPublicOnly)) {')
add_code('        wrapper.eq(Message::getIsPublic, true);')
add_code('    } else if (currentUserId != null) {')
add_code('        wrapper.and(w -> w.eq(Message::getIsPublic, true)')
add_code('                .or().eq(Message::getUserId, currentUserId));')
add_code('    } else {')
add_code('        wrapper.eq(Message::getIsPublic, true);')
add_code('    }')
add_code('    // ... 其他筛选条件 ...')
add_code('    return page(new Page<>(page, size), wrapper);')
add_code('}')

add_title('5.4 前端实现', level=2)

add_body('前端API层封装了11个省长信箱相关接口函数：submitMessage、messageList、messageDetail、hotMessages、replyMessage、userReplyMessage、dispatchMessage、finishMessage、superviseMessage、rateMessage、messageStats。')

add_body('C端页面包括WriteLetter.vue（写信页面，邮件式界面，包含收件人选择、主题输入、富文本工具栏、HTML正文编辑区）、LetterDetail.vue（信件详情页，展示完整信息和多轮对话记录，区分管理员/群众消息样式，仅本人显示追问输入框）、PublicLetters.vue（公开信件列表）、Interaction.vue（政民互动首页，含热门信件展示和评价入口）。')

add_body('B端页面为InteractionMgmt.vue（省长信箱管理，包含信件列表、分派弹窗、回复弹窗、督办操作、办结操作、统计面板）。')

doc.add_page_break()

# ============================================================
# 六、意见征集
# ============================================================
add_title('六、意见征集功能开发与实现', level=1)

add_title('6.1 功能描述', level=2)

add_body('意见征集功能实现了"政府发布→群众提交→管理员处理→结果公示"的完整闭环。与咨询/建议/投诉的单向问答模式不同，意见征集是双向流程：政府先发布征集活动，群众在征集期间提交意见，管理员逐条处理意见（采纳/部分采纳/不采纳），最后发布结果反馈公示。')

add_body('该功能涉及2张数据表：collection（征集主表）和collection_opinion（意见提交表），通过collection_id建立一对多关联关系。')

add_body('该功能包含以下13个功能点：')

add_bullet('征集列表（C端）：群众分页查看公开征集活动列表，支持关键词和状态筛选。')
add_bullet('征集详情（C端）：群众查看征集主题、说明、征集单位、联系人、时间范围、附件，以及意见统计（总数/待处理/已回复）。')
add_bullet('提交意见（C端）：群众填写意见标题、姓名、手机号码、身份证号、通信地址、留言内容。后端校验征集存在性和必填字段后保存。')
add_bullet('结果反馈公示（C端）：群众查看征集结果反馈内容和意见采纳统计（已采纳/部分采纳/不采纳数量）。')
add_bullet('创建征集（B端）：管理员填写征集主题、说明、征集单位、联系人、联系电话、附件、时间范围。系统根据当前日期与开始/结束日期的关系自动设置状态（未开始/进行中/已结束）。')
add_bullet('编辑征集（B端）：修改征集信息，校验开始日期不晚于结束日期。')
add_bullet('删除征集（B端）：逻辑删除征集活动。')
add_bullet('发布征集（B端）：将征集状态设为"进行中"，开放群众提交意见。')
add_bullet('结束征集（B端）：将征集状态设为"已结束"，停止接收新意见。')
add_bullet('结果反馈（B端）：添加反馈内容，记录反馈时间，用于C端公示。')
add_bullet('意见列表（B端）：查看某征集下的所有意见，支持关键词和状态筛选。')
add_bullet('回复意见（B端）：处理意见，选择处理状态（已采纳/部分采纳/不采纳），填写反馈理由。后端白名单校验状态值。')
add_bullet('意见统计（B端）：统计某征集的意见数量，包括总数、待处理、已回复、已采纳、部分采纳、不采纳6个维度。')

add_title('6.2 业务流程', level=2)

add_body('意见征集的完整业务流程如下：')

add_bullet('管理员在后台创建征集活动，填写主题、说明、单位、联系人、时间范围等信息。')
add_bullet('系统根据当前日期自动判断状态：当前日期在开始日期之前为"未开始"，在时间范围内为"进行中"，在结束日期之后为"已结束"。')
add_bullet('管理员发布征集，状态变为"进行中"，C端群众可以查看并提交意见。')
add_bullet('群众在C端查看征集列表，点击感兴趣的征集查看详情。')
add_bullet('群众填写意见表单（标题/姓名/手机/身份证/地址/内容），提交后状态为"待处理"。')
add_bullet('管理员在后台查看意见列表，逐条处理意见。')
add_bullet('管理员选择处理结果（已采纳/部分采纳/不采纳），填写反馈理由，后端校验状态值合法性后更新。')
add_bullet('征集结束后，管理员添加结果反馈内容。')
add_bullet('C端群众可在结果反馈公示页面查看反馈内容和采纳统计。')

add_title('6.3 核心代码', level=2)

add_body('创建征集方法（自动状态判断）：', bold=True)
add_code('public Collection create(Collection collection) {')
add_code('    if (!StringUtils.hasText(collection.getTitle()))')
add_code('        throw BusinessException.of("征集主题不能为空");')
add_code('    if (collection.getStartDate() == null || collection.getEndDate() == null)')
add_code('        throw BusinessException.of("征集时间不能为空");')
add_code('    if (collection.getStartDate().isAfter(collection.getEndDate()))')
add_code('        throw BusinessException.of("开始日期不能晚于结束日期");')
add_code('    // 根据日期自动设置状态')
add_code('    LocalDate today = LocalDate.now();')
add_code('    if (today.isBefore(collection.getStartDate())) {')
add_code('        collection.setStatus("未开始");')
add_code('    } else if (today.isAfter(collection.getEndDate())) {')
add_code('        collection.setStatus("已结束");')
add_code('    } else {')
add_code('        collection.setStatus("进行中");')
add_code('    }')
add_code('    save(collection);')
add_code('    return collection;')
add_code('}')

add_body('提交意见方法：', bold=True)
add_code('public CollectionOpinion submitOpinion(CollectionOpinion opinion) {')
add_code('    if (opinion.getCollectionId() == null)')
add_code('        throw BusinessException.of("征集ID不能为空");')
add_code('    Collection c = getById(opinion.getCollectionId());')
add_code('    if (c == null)')
add_code('        throw BusinessException.of("征集不存在");')
add_code('    if (!StringUtils.hasText(opinion.getTitle()))')
add_code('        throw BusinessException.of("意见标题不能为空");')
add_code('    if (!StringUtils.hasText(opinion.getRealName()))')
add_code('        throw BusinessException.of("姓名不能为空");')
add_code('    if (!StringUtils.hasText(opinion.getPhone()))')
add_code('        throw BusinessException.of("手机号码不能为空");')
add_code('    if (!StringUtils.hasText(opinion.getContent()))')
add_code('        throw BusinessException.of("留言内容不能为空");')
add_code('    opinion.setStatus("待处理");')
add_code('    opinionMapper.insert(opinion);')
add_code('    return opinion;')
add_code('}')

add_body('回复意见方法（采纳/不采纳，白名单校验）：', bold=True)
add_code('public void replyOpinion(Long opinionId, String replyContent,')
add_code('        String operator, String status) {')
add_code('    CollectionOpinion opinion = opinionMapper.selectById(opinionId);')
add_code('    if (opinion == null)')
add_code('        throw BusinessException.of("意见不存在");')
add_code('    if (!StringUtils.hasText(replyContent))')
add_code('        throw BusinessException.of("反馈理由不能为空");')
add_code('    if (!"已采纳".equals(status) && !"部分采纳".equals(status)')
add_code('            && !"不采纳".equals(status))')
add_code('        throw BusinessException.of("处理状态只能是已采纳、部分采纳或不采纳");')
add_code('    opinion.setReplyContent(replyContent);')
add_code('    opinion.setReplyBy(operator);')
add_code('    opinion.setReplyTime(LocalDateTime.now());')
add_code('    opinion.setStatus(status);')
add_code('    opinionMapper.updateById(opinion);')
add_code('}')

add_body('意见统计方法（6个维度）：', bold=True)
add_code('public Map<String, Long> opinionStats(Long collectionId) {')
add_code('    Map<String, Long> stats = new HashMap<>();')
add_code('    LambdaQueryWrapper<CollectionOpinion> base = new LambdaQueryWrapper<>();')
add_code('    base.eq(CollectionOpinion::getCollectionId, collectionId);')
add_code('    stats.put("total", opinionMapper.selectCount(base));')
add_code('    // 待处理')
add_code('    LambdaQueryWrapper<CollectionOpinion> pw = new LambdaQueryWrapper<>();')
add_code('    pw.eq(CollectionOpinion::getCollectionId, collectionId)')
add_code('      .eq(CollectionOpinion::getStatus, "待处理");')
add_code('    stats.put("pending", opinionMapper.selectCount(pw));')
add_code('    // 已回复（已采纳+部分采纳+不采纳）')
add_code('    LambdaQueryWrapper<CollectionOpinion> rw = new LambdaQueryWrapper<>();')
add_code('    rw.eq(CollectionOpinion::getCollectionId, collectionId)')
add_code('      .in(CollectionOpinion::getStatus, "已采纳","部分采纳","不采纳");')
add_code('    stats.put("replied", opinionMapper.selectCount(rw));')
add_code('    // 已采纳')
add_code('    LambdaQueryWrapper<CollectionOpinion> aw = new LambdaQueryWrapper<>();')
add_code('    aw.eq(CollectionOpinion::getCollectionId, collectionId)')
add_code('      .eq(CollectionOpinion::getStatus, "已采纳");')
add_code('    stats.put("adopted", opinionMapper.selectCount(aw));')
add_code('    // 部分采纳')
add_code('    LambdaQueryWrapper<CollectionOpinion> paw = new LambdaQueryWrapper<>();')
add_code('    paw.eq(CollectionOpinion::getCollectionId, collectionId)')
add_code('       .eq(CollectionOpinion::getStatus, "部分采纳");')
add_code('    stats.put("partiallyAdopted", opinionMapper.selectCount(paw));')
add_code('    // 不采纳')
add_code('    LambdaQueryWrapper<CollectionOpinion> rjw = new LambdaQueryWrapper<>();')
add_code('    rjw.eq(CollectionOpinion::getCollectionId, collectionId)')
add_code('       .eq(CollectionOpinion::getStatus, "不采纳");')
add_code('    stats.put("rejected", opinionMapper.selectCount(rjw));')
add_code('    return stats;')
add_code('}')

add_title('6.4 前端实现', level=2)

add_body('前端API层封装了14个意见征集相关接口函数，分为C端和B端两组。C端4个：collectionPublicList、collectionPublicDetail、collectionPublicFeedback、submitCollectionOpinion。B端10个：collectionList、createCollection、updateCollection、deleteCollection、publishCollection、finishCollection、addCollectionFeedback、collectionOpinions、replyCollectionOpinion、collectionOpinionStats。')

add_body('C端页面包括CollectionList.vue（征集列表）、CollectionDetail.vue（征集详情+意见统计）、SubmitOpinion.vue（提交意见表单）、CollectionFeedback.vue（结果反馈公示页，展示采纳统计）。')

add_body('B端页面为CollectionMgmt.vue（意见征集管理，包含征集CRUD、发布/结束操作、意见列表、意见处理弹窗、反馈编辑、统计面板，是B端功能最丰富的管理页面）。')

doc.add_page_break()

# ============================================================
# 七、接口汇总
# ============================================================
add_title('七、接口汇总', level=1)

add_body('政民互动板块共开发46个RESTful接口，分布在5个功能模块中。以下按功能列出所有接口：')

add_body('我要咨询（7个接口，前缀 /api/v1/consultation）：', bold=True)
add_bullet('POST / — 提交咨询')
add_bullet('GET /progress/{consultNo} — 查询进度（仅本人）')
add_bullet('GET /list — 咨询列表（分页+筛选）')
add_bullet('POST /{consultNo}/reply — 答复咨询')
add_bullet('PATCH /{consultNo}/status — 修改状态')
add_bullet('POST /{consultNo}/finish — 办结咨询')
add_bullet('GET /stats — 数据统计')

add_body('我要建议（7个接口，前缀 /api/v1/suggestion）：', bold=True)
add_bullet('POST / — 提交建议')
add_bullet('GET /progress/{suggestNo} — 查询进度（仅本人）')
add_bullet('GET /list — 建议列表（分页+筛选）')
add_bullet('POST /{suggestNo}/reply — 答复建议')
add_bullet('PATCH /{suggestNo}/status — 修改状态')
add_bullet('POST /{suggestNo}/finish — 办结建议')
add_bullet('GET /stats — 数据统计')

add_body('我要投诉（7个接口，前缀 /api/v1/complaint）：', bold=True)
add_bullet('POST / — 提交投诉')
add_bullet('GET /progress/{complaintNo} — 查询进度（仅本人）')
add_bullet('GET /list — 投诉列表（分页+筛选）')
add_bullet('POST /{complaintNo}/reply — 答复投诉')
add_bullet('PATCH /{complaintNo}/status — 修改状态')
add_bullet('POST /{complaintNo}/finish — 办结投诉')
add_bullet('GET /stats — 数据统计')

add_body('省长信箱（11个接口，前缀 /api/v1/interaction/message）：', bold=True)
add_bullet('POST / — 提交信件')
add_bullet('GET / — 信件列表（C端可见性过滤/B端全部）')
add_bullet('GET /{id} — 信件详情（含对话记录+可见性校验）')
add_bullet('GET /hot — 热门信件')
add_bullet('POST /{id}/reply — 管理员回复')
add_bullet('POST /{id}/user-reply — 群众追问（仅本人）')
add_bullet('POST /{id}/dispatch — 分派管理')
add_bullet('POST /{id}/supervise — 督办管理')
add_bullet('POST /{id}/finish — 办结信件')
add_bullet('POST /{id}/rate — 满意度评价')
add_bullet('GET /stats — 数据统计')

add_body('意见征集（14个接口，前缀 /api/v1/collection）：', bold=True)
add_bullet('GET /public — C端征集列表')
add_bullet('GET /public/{id} — C端征集详情')
add_bullet('GET /public/{id}/feedback — C端结果反馈公示')
add_bullet('POST /{id}/opinion — C端提交意见')
add_bullet('GET / — B端征集列表')
add_bullet('POST / — B端创建征集')
add_bullet('PUT /{id} — B端更新征集')
add_bullet('DELETE /{id} — B端删除征集')
add_bullet('POST /{id}/publish — B端发布征集')
add_bullet('POST /{id}/finish — B端结束征集')
add_bullet('POST /{id}/feedback — B端添加结果反馈')
add_bullet('GET /{id}/opinions — B端意见列表')
add_bullet('POST /opinion/{opinionId}/reply — B端回复意见')
add_bullet('GET /{id}/opinion-stats — B端意见统计')

doc.add_page_break()

# ============================================================
# 八、实现截图说明
# ============================================================
add_title('八、实现截图说明', level=1)

add_body('以下为各功能的实现截图位置说明，请在系统运行后截取对应页面插入文档：')

add_body('我要咨询截图：', bold=True)
add_bullet('C端咨询提交页面（路由 /interaction/consultation）：展示咨询表单，包含姓名、电话、主题、市州选择、内容编辑区、附件上传、公开/保密选项')
add_bullet('C端咨询进度查询页面：输入单号后展示处理状态和答复内容')
add_bullet('B端咨询管理列表页面（路由 /admin/consultation）：展示咨询列表、筛选栏、答复弹窗、状态修改、办结按钮、统计面板')

add_body('我要建议截图：', bold=True)
add_bullet('C端建议提交页面（路由 /interaction/suggest）：展示建议表单，包含建议类型下拉框、三级地址选择器')
add_bullet('C端建议详情页面：展示建议信息和答复内容')
add_bullet('B端建议管理列表页面（路由 /admin/suggestion）：功能布局与咨询管理一致')

add_body('我要投诉截图：', bold=True)
add_bullet('C端投诉提交页面（路由 /interaction/complaint）：展示投诉表单，包含三级投诉单位联动选择器、三级问题类型联动选择器')
add_bullet('C端投诉进度查询页面')
add_bullet('B端投诉管理列表页面（路由 /admin/complaint）')

add_body('省长信箱截图：', bold=True)
add_bullet('C端写信页面（路由 /interaction/write）：邮件式界面，包含收件人选择、主题输入、富文本工具栏、HTML正文编辑区')
add_bullet('C端公开信件列表页面')
add_bullet('C端信件详情页面（路由 /interaction/letter/:id）：展示完整信息和多轮对话记录，区分管理员/群众消息样式')
add_bullet('C端群众追问输入区域（仅本人可见）')
add_bullet('C端满意度评价弹窗')
add_bullet('B端省长信箱管理页面（路由 /admin/interaction）：包含分派、督办、回复、办结操作')

add_body('意见征集截图：', bold=True)
add_bullet('C端意见征集列表页面（路由 /interaction/collection）')
add_bullet('C端征集详情页面：展示征集说明和意见统计')
add_bullet('C端提交意见页面')
add_bullet('C端结果反馈公示页面：展示采纳统计')
add_bullet('B端意见征集管理页面（路由 /admin/collection）：包含征集CRUD、发布/结束、意见列表、意见处理、反馈编辑')

# ============================================================
# 保存文件
# ============================================================
output_path = r'E:\huadee\政民互动功能开发与实现.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
print(f'文件大小：{os.path.getsize(output_path) / 1024:.1f} KB')
