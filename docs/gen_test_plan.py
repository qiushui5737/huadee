# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ========== 全局样式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def set_cell_font(cell, text, bold=False, size=Pt(10.5), font_name='宋体', alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return h

def add_para(doc, text, bold=False, size=Pt(12), alignment=None, indent=False):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, bold=True, size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_font(table.rows[r_idx + 1].cells[c_idx], str(val), size=Pt(10.5))
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ========== 封面 ==========
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('项目编号：001')
run.font.name = '宋体'
run.font.size = Pt(14)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('政府网站集约化平台')
run.font.name = '黑体'
run.font.size = Pt(26)
run.font.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('测试计划')
run.font.name = '黑体'
run.font.size = Pt(22)
run.font.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0.1')
run.font.name = '宋体'
run.font.size = Pt(14)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for _ in range(3):
    doc.add_paragraph()

cover_info = [
    ('项 目 承 担 部 门：', '4组'),
    ('撰  写  人（签名）：', '蒙焕好'),
    ('完 成 日 期：', '2026-7-19'),
]
for label, val in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + '  ' + val)
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('本文档使用部门： ■主管领导  ■项目组  □客户（市场）  □维护人员  □用户')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('评审负责人（签名）：叶诚        评审日期：2026-7-20')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ========== 文档信息 ==========
add_heading_styled(doc, '文档信息', 1)
make_table(doc,
    ['标题', '遂宁市政府网站集约化平台-测试计划'],
    [['作者', '蒙焕好'], ['创建日期', '2026年7月16日'], ['版本', 'V1.0.1'], ['部门名称', '4组']]
)
doc.add_paragraph()
add_para(doc, '修订文档历史记录', bold=True)
make_table(doc,
    ['日期', '版本', '说明', '作者'],
    [['2026/7/16', '1.0', '草稿', '蒙焕好'], ['2026/7/19', '1.0.1', '正式发布', '蒙焕好']]
)

doc.add_page_break()

# ========== 目录占位 ==========
add_para(doc, '目  录', bold=True, size=Pt(16), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '（请在Word中插入自动目录：引用 → 目录 → 自动目录）', size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ========== 1. 简介 ==========
add_heading_styled(doc, '1. 简介', 1)

add_heading_styled(doc, '1.1 目的', 2)
add_para(doc, '遂宁市政府网站集约化平台"软件测试计划"文档有助于实现以下目标：', indent=True)
for t in ['确定遂宁市政府网站集约化平台的信息和应测试的软件构件。',
          '确定测试需求。',
          '确定测试策略，并对这些策略加以说明。',
          '确定所需的资源，并对测试的工作量进行估计。',
          '列出项目测试的可交付工件。']:
    add_bullet(doc, t)

add_heading_styled(doc, '1.2 背景', 2)
add_para(doc, '遂宁市政府网站集约化平台是面向公众提供政务服务、政民互动、信息公开等功能的综合性政府门户网站。平台分为客户端（C端）和管理端（B端），客户端面向广大市民提供办事服务、在线咨询、投诉、建议、意见征集、依申请公开、AI智能问答等功能；管理端面向政府工作人员提供互动管理、内容审核、服务管理、绩效统计、AI审计等后台管理功能。', indent=True)
add_para(doc, '平台采用前后端分离架构，后端基于SpringBoot 3.2.5 + MyBatis-Plus + MySQL + Redis，前端基于Vue3 + TypeScript + Element Plus + Vite，包含以下核心子系统：', indent=True)
for t in ['政民互动子系统：留言系统、咨询、投诉、建议、意见征集、依申请公开、民意调查',
          '办事服务子系统：事项目录管理、在线办理',
          '内容管理子系统：内容发布、审核、信息公开',
          'AI智能服务子系统：智能问答、语义搜索、内容审计',
          '系统管理子系统：用户认证、权限管理、绩效统计']:
    add_bullet(doc, t)

add_heading_styled(doc, '1.3 范围', 2)
add_para(doc, '对于本项目的测试工作为系统测试阶段，阶段的测试活动有计划测试、设计测试、执行测试和评估测试四个活动。', indent=True)
add_para(doc, '本测试计划仅对项目开发的软件提供测试，不考虑对第三方提供的软件（如第三方的组件或构件）进行测试。但需测试第三方提供的软件与开发的软件之间的接口。', indent=True)
add_para(doc, '本测试计划适用于"遂宁市政府网站集约化平台"项目。本文档将供给项目经理及项目开发各组使用，包括测试组、开发组。', indent=True)

add_heading_styled(doc, '1.4 定义', 2)
add_para(doc, '缺陷级别定义：', indent=True)
for t in ['一级：不能完全满足系统要求，基本功能未完全实现；或者危及人身安全。',
          '二级：严重地影响系统要求或基本功能的实现，且没有更正办法（重新安装或重新启动该软件不属于更正办法）。',
          '三级：严重地影响系统要求或基本功能的实现，但存在合理的更正办法（重新安装或重新启动该软件不属于更正办法）。',
          '四级：使操作者不方便或遇到麻烦，但它不影响执行工作功能或重要功能。',
          '五级：其他缺陷。']:
    add_bullet(doc, t)

add_heading_styled(doc, '1.5 文档引用', 2)
add_bullet(doc, '《遂宁市政府网站集约化平台开发计划》')
add_bullet(doc, '《遂宁市政府网站集约化平台需求规格说明书》')

# ========== 2. 测试需求 ==========
add_heading_styled(doc, '2. 测试需求', 1)
add_para(doc, '备注：优先级：1-高，2-中，3-低', size=Pt(10.5))

add_heading_styled(doc, '2.1 功能测试需求', 2)
func_rows = [
    ['用户注册与登录', 'C端用户通过用户名、密码、实名信息进行注册；B端管理员通过账户、密码方式登录系统，获取用户信息', '1'],
    ['政民互动留言', 'C端群众提交留言（含咨询、投诉、建议、求助类型），查看留言列表与回复，进行评价；B端管理员分派、回复、督办留言', '1'],
    ['我要咨询', 'C端用户提交咨询问题，生成咨询单号，根据单号查询咨询进度；B端管理员查看咨询列表、回复咨询', '1'],
    ['我要投诉', 'C端用户填写投诉信息提交投诉，生成投诉单号，查询投诉进度；B端管理员查看投诉列表、回复处理', '1'],
    ['我要建议', 'C端用户提交建议内容，生成建议单号，查询建议进度；B端管理员查看建议列表、回复处理', '1'],
    ['意见征集', 'C端用户查看征集列表与详情，提交意见；B端管理员新建/编辑/发布/停止征集，对意见进行采纳/部分采纳/不采纳处理，填写结果反馈', '1'],
    ['依申请公开', 'C端群众提交政府信息公开申请；B端管理员受理、审核、答复、归档', '1'],
    ['民意调查问卷', 'B端管理员创建问卷（单选/多选/文本题型），发布/关闭问卷；C端用户填写并提交问卷', '2'],
    ['办事服务管理', 'B端管理员管理办事事项（新增/编辑/删除/查询）；C端用户浏览事项目录、查看事项详情', '1'],
    ['内容管理与信息公开', 'B端管理员发布内容并进入审核流程；C端用户查看公开信息', '2'],
    ['AI智能问答', 'C端用户通过AI对话进行政务咨询，支持流式回答、会话历史；B端管理员查看AI审计记录', '2'],
    ['AI智能搜索', 'C端用户通过关键词搜索政务信息，支持语义搜索和热门搜索', '2'],
    ['绩效统计', 'B端管理员查看仪表盘核心指标及周趋势数据，生成统计报表', '2'],
    ['平台人员管理', 'B端管理员对平台用户进行管理，包括注册审核、角色分配', '1'],
    ['退出登录', '后台管理员登录后退出系统，清除登录会话', '3'],
]
make_table(doc, ['测试需求', '测试需求简介', '优先级'], func_rows, col_widths=[3.5, 10, 1.5])

add_heading_styled(doc, '2.2 非功能测试需求', 2)
nfunc_rows = [
    ['性能-查询响应时间', '列表查询、详情查询等常规查询操作', '查询结果响应时间为3秒以内'],
    ['性能-页面响应时间', '页面加载、路由切换、表单提交等操作', '页面响应时间1秒以内'],
    ['性能-并发处理', '多用户同时访问系统', '系统支持至少50个并发用户正常使用'],
    ['界面需求', '前端界面', '采用现代化Web图形界面，支持鼠标键盘操作；对用户友好，能正常运行'],
    ['兼容性', '浏览器兼容', '支持Chrome、Firefox、Edge、360等主流浏览器，界面正常显示，功能运行正常'],
    ['安全性-认证与授权', 'JWT Token认证', 'C端/B端权限隔离，未登录用户无法访问受保护接口'],
    ['安全性-数据保护', '密码与敏感信息', '用户密码加密存储（BCrypt），敏感信息脱敏处理'],
    ['可靠性', '异常处理', '网络异常、数据校验失败等场景有明确错误提示，系统不会崩溃'],
]
make_table(doc, ['测试需求项', '条件', '指标'], nfunc_rows, col_widths=[3.5, 5, 6.5])

# ========== 3. 测试策略 ==========
add_heading_styled(doc, '3. 测试策略', 1)
add_heading_styled(doc, '3.1 测试类型', 2)
add_para(doc, '本次测试为系统测试。具体如下：', indent=True)

add_heading_styled(doc, '3.1.1 功能测试策略', 3)
add_para(doc, '测试目标', bold=True)
add_para(doc, '所有功能测试需求项的功能实现。', indent=True)
add_para(doc, '测试方法和技术', bold=True)
add_para(doc, '按照测试需求、通过准则、测试用例，采用黑盒法测试，核实以下内容：', indent=True)
for t in ['在使用合法数据时得到正确的结果。',
          '在使用非法数据时显示相应的缺陷信息或警告信息。',
          '各业务规则都得到了正确的应用。',
          'C端与B端数据交互正确（如C端提交→B端可见，B端回复→C端可查）。',
          '分页查询、条件筛选、排序等功能正确。',
          '表单校验（必填项、格式校验、长度限制等）生效。']:
    add_bullet(doc, t)
add_para(doc, '完成标准', bold=True)
for t in ['所计划的测试已全部执行。', '所发现的缺陷修复率达到通过准则要求。', '提供《测试记录》。']:
    add_bullet(doc, t)
add_para(doc, '需考虑的特殊事项', bold=True)
for t in ['意见征集的三态处理（已采纳/部分采纳/不采纳）需重点验证。',
          '政民互动留言的分派督办流程需验证完整闭环。',
          '依申请公开的时限管理需关注截止日期计算。',
          'JWT Token过期和会话管理需验证Redis中Token的一致性。']:
    add_bullet(doc, t)

add_heading_styled(doc, '3.1.2 性能测试策略', 3)
add_para(doc, '测试目标', bold=True)
add_para(doc, '所有性能测试需求项的正常实现。', indent=True)
add_para(doc, '测试方法和技术', bold=True)
for t in ['测试环境参见"系统资源"。',
          '主要针对数据处理能力、响应时间做测试。',
          '对列表查询接口进行分页压力测试（1000条以上数据量）。',
          '对AI流式问答接口进行响应时间测试。',
          '具体计划在测试时制定。']:
    add_bullet(doc, t)
add_para(doc, '完成标准', bold=True)
for t in ['所计划的测试已全部执行。', '系统运行稳定。', '所有性能测试需求项符合要求。', '提供《系统性能测试记录表》。']:
    add_bullet(doc, t)
add_para(doc, '需考虑的特殊事项', bold=True)
add_para(doc, '无', indent=True)

add_heading_styled(doc, '3.1.3 兼容性测试策略', 3)
add_para(doc, '测试目标', bold=True)
add_para(doc, '支持Chrome、Firefox、Edge、360等主流浏览器。', indent=True)
add_para(doc, '测试方法和技术', bold=True)
add_para(doc, '在浏览器下运行系统，验证页面布局、交互功能、表单操作等是否正常。', indent=True)
add_para(doc, '完成标准', bold=True)
add_para(doc, '各种浏览器下正常显示，功能正常运行。', indent=True)
add_para(doc, '需考虑的特殊事项', bold=True)
for t in ['需特别关注富文本编辑器在Edge浏览器中的兼容性问题。',
          '需验证Element Plus组件库在各浏览器中的渲染一致性。']:
    add_bullet(doc, t)

add_heading_styled(doc, '3.2 测试工具', 2)
make_table(doc,
    ['名称', '工具', '版本'],
    [['功能性测试工具', 'QC', '9.0'],
     ['性能测试工具', 'LoadRunner', '11.0'],
     ['项目管理', 'Ms Project', '2010'],
     ['接口测试工具', 'Postman', '最新版'],
     ['浏览器调试', 'Chrome DevTools', '最新版']])

add_heading_styled(doc, '3.3 测试通过准则', 2)
for t in ['实行了所有的测试策略并达到完成标准。',
          '测试结束后，开发组对实现有误的测试需求项的修改达到如下程度：一、二级缺陷的修复率达到100%；三、四级缺陷的修复率达到80%以上；五级缺陷的修复率达到60%以上。',
          '《需求规格说明书》、《用户手册》、《操作手册》和编码实现一致。']:
    add_bullet(doc, t)

# ========== 4. 资源 ==========
add_heading_styled(doc, '4. 资源', 1)
add_heading_styled(doc, '4.1 角色', 2)
make_table(doc,
    ['人力资源', '角色', '具体职责或注释'],
    [['杨雨洁', '测试经理', '进行管理监督。职责：提供技术指导、获取适当的资源、提供管理报告'],
     ['邬中塬、付杰', '测试设计员', '计划测试、设计测试及评估测试。职责：确定测试用例、确定测试用例的优先级并实施测试用例、生成测试计划、生成测试模型、评估测试工作的有效性'],
     ['李欣、杨坤荣', '测试员', '执行测试。职责：执行测试、记录结果、从缺陷中恢复']])

add_heading_styled(doc, '4.2 系统资源', 2)
add_para(doc, '硬件环境：', bold=True)
make_table(doc, ['环境', '配置'], [['服务端', '由客户提供'], ['客户端测试机', '普通PC，内存8GB以上']])
doc.add_paragraph()
add_para(doc, '软件环境：', bold=True)
make_table(doc, ['类别', '配置'],
    [['数据库', 'MySQL 8.x'],
     ['缓存', 'Redis'],
     ['后端框架', 'SpringBoot 3.2.5 + MyBatis-Plus 3.5.5'],
     ['前端框架', 'Vue3 + TypeScript + Element Plus + Vite 5'],
     ['运行环境', 'Java 17、Node.js 18+'],
     ['操作系统', 'Windows 10/11']])

# ========== 5. 项目测试活动 ==========
add_heading_styled(doc, '5. 项目测试活动', 1)
make_table(doc,
    ['里程碑任务', '工作', '开始日期', '结束日期', '交付工件'],
    [['计划测试', '阅读参考资料，制定测试计划，分派测试任务。', '2025.8.18', '2025.8.20', '《测试计划》'],
     ['设计测试', '根据测试计划、用例阐述和设计来设计测试用例，可在测试过程中进行。', '2025.8.20', '2025.8.21', '《测试用例》'],
     ['执行测试', '进行系统测试，包括功能测试、性能测试、兼容测试', '2025.8.22', '2025.8.23', '《测试记录》'],
     ['评估测试', '编写《测试分析报告》。', '2025.8.23', '2025.8.25', '《测试分析报告》']])

# ========== 6. 附录A ==========
add_heading_styled(doc, '6. 附录 A：项目任务', 1)
add_para(doc, '以下是一些与测试有关的任务：', indent=True)

for section_title, items in [
    ('制定测试计划', ['确定测试需求', '评估风险', '制定测试策略', '确定测试资源', '创建时间表', '生成测试计划']),
    ('设计测试', ['准备工作量分析文档', '确定并说明测试用例', '确定测试过程，并建立测试过程的结构', '复审和评估测试覆盖']),
    ('执行测试', ['执行测试过程', '评估测试的执行情况', '恢复暂停的测试', '核实结果', '调查意外结果', '记录缺陷']),
    ('对测试进行评估', ['评估测试用例覆盖', '评估代码覆盖', '分析缺陷', '确定是否达到了测试完成标准与成功标准']),
]:
    add_para(doc, section_title, bold=True)
    for item in items:
        add_bullet(doc, item)

# ========== 附录B 测试用例对照表 ==========
doc.add_page_break()
add_heading_styled(doc, '附录 B：功能模块与测试用例对照表', 1)
add_para(doc, '以下为各功能模块的详细测试要点，供测试设计员编写测试用例时参考：', indent=True)

test_cases = [
    ('B1 用户注册与登录', [
        ['T01-01', 'C端用户使用合法用户名、密码、实名信息注册', '注册成功，提示"注册成功"'],
        ['T01-02', 'C端用户注册时用户名为空', '提示"请输入用户名"'],
        ['T01-03', 'C端用户注册时密码为空', '提示"请输入密码"'],
        ['T01-04', 'C端用户注册时真实姓名为空', '提示"请输入真实姓名"'],
        ['T01-05', 'C端用户使用已存在的用户名注册', '提示用户名已存在'],
        ['T01-06', 'B端管理员使用正确账号密码登录', '登录成功，跳转管理后台首页'],
        ['T01-07', 'B端管理员使用错误密码登录', '提示"用户名或密码错误"'],
        ['T01-08', '登录后调用/auth/info获取用户信息', '返回用户ID、用户名、真实姓名、角色等信息'],
        ['T01-09', 'Token过期后访问受保护接口', '返回401"未登录或登录已过期"'],
        ['T01-10', 'C端用户访问B端管理接口', '返回403"无管理端访问权限"'],
        ['T01-11', '退出登录', '会话清除，再次访问需重新登录'],
    ]),
    ('B2 政民互动留言', [
        ['T02-01', 'C端用户提交留言（填写标题、内容、联系人、目标部门）', '留言提交成功，返回留言ID'],
        ['T02-02', 'C端用户提交留言时标题为空', '表单校验不通过，提示必填'],
        ['T02-03', 'C端用户查看留言列表', '分页展示留言列表，显示标题、状态、回复等'],
        ['T02-04', 'C端用户查看留言详情及回复', '显示留言完整内容及管理员回复'],
        ['T02-05', 'C端用户对已回复留言进行评价', '评价提交成功'],
        ['T02-06', 'B端管理员查看留言列表', '分页展示所有留言，支持关键词搜索'],
        ['T02-07', 'B端管理员回复留言', '回复成功，C端可查看到回复内容'],
        ['T02-08', 'B端管理员分派留言至指定部门', '分派成功，留言状态更新'],
    ]),
    ('B3 我要咨询', [
        ['T03-01', 'C端用户提交咨询（姓名、电话、内容等）', '提交成功，生成咨询单号'],
        ['T03-02', 'C端用户根据咨询单号查询进度', '返回咨询详情及当前状态'],
        ['T03-03', 'C端用户查询非本人咨询单号', '返回无权限或查无数据'],
        ['T03-04', 'B端管理员查看咨询列表', '分页展示，支持关键词和状态筛选'],
        ['T03-05', 'B端管理员回复咨询', '回复成功，状态更新为"已回复"'],
    ]),
    ('B4 我要投诉', [
        ['T04-01', 'C端用户提交投诉（投诉单位、问题类型、个人信息、内容等）', '提交成功，生成投诉单号及截止日期'],
        ['T04-02', 'C端用户查询本人投诉进度', '返回投诉详情、回复及处理状态'],
        ['T04-03', 'C端用户查询非本人投诉单号', '返回无权限'],
        ['T04-04', 'B端管理员查看投诉列表', '分页展示，支持关键词、状态筛选'],
        ['T04-05', 'B端管理员回复投诉', '回复成功'],
        ['T04-06', '投诉信息中设置公开/保密选项', '公开信件在C端可见，保密信件仅本人可见'],
    ]),
    ('B5 我要建议', [
        ['T05-01', 'C端用户提交建议（标题、内容、联系方式等）', '提交成功，生成建议单号'],
        ['T05-02', 'C端用户根据建议单号查询进度', '返回建议详情及处理状态'],
        ['T05-03', 'B端管理员查看建议列表', '分页展示，支持筛选'],
        ['T05-04', 'B端管理员回复建议', '回复成功'],
    ]),
    ('B6 意见征集', [
        ['T06-01', 'B端管理员新建征集（主题、说明、单位、时间等）', '创建成功'],
        ['T06-02', 'B端管理员编辑已创建的征集', '编辑保存成功'],
        ['T06-03', 'B端管理员发布征集', '状态变为"进行中"，C端可见'],
        ['T06-04', 'B端管理员停止进行中的征集', '状态变为"已结束"'],
        ['T06-05', 'C端用户查看征集列表', '分页展示，支持关键词搜索和状态筛选'],
        ['T06-06', 'C端用户查看征集详情', '显示征集主题、时间、单位、说明材料等'],
        ['T06-07', 'C端用户在"进行中"征集提交意见', '意见提交成功'],
        ['T06-08', 'C端用户在"已结束"征集中尝试提交意见', '无法提交，提示征集已结束'],
        ['T06-09', 'B端管理员对意见标记"已采纳"', '意见状态更新为"已采纳"'],
        ['T06-10', 'B端管理员对意见标记"部分采纳"并填写理由', '意见状态更新，理由保存成功'],
        ['T06-11', 'B端管理员对意见标记"不采纳"并填写理由', '意见状态更新，理由保存成功'],
        ['T06-12', 'B端管理员查看意见统计（总数/已采纳/部分采纳/不采纳）', '统计数据正确'],
        ['T06-13', 'B端管理员填写结果反馈内容', '反馈内容保存成功'],
        ['T06-14', 'C端用户查看结果反馈公示页面', '显示反馈内容、意见统计图表'],
    ]),
    ('B7 依申请公开', [
        ['T07-01', 'C端用户提交公开申请（申请人、身份证、电话、内容、获取方式）', '申请提交成功'],
        ['T07-02', 'C端用户查看申请列表', '分页展示申请记录'],
        ['T07-03', 'B端管理员查看申请列表', '分页展示，支持筛选'],
        ['T07-04', 'B端管理员受理申请', '状态更新为"已受理"'],
        ['T07-05', 'B端管理员答复申请', '答复内容保存，状态更新'],
        ['T07-06', 'B端管理员归档已答复申请', '归档成功'],
    ]),
    ('B8 民意调查问卷', [
        ['T08-01', 'B端管理员创建问卷（标题、题目、题型）', '创建成功'],
        ['T08-02', 'B端管理员发布问卷', '状态变为"已发布"，C端可见'],
        ['T08-03', 'C端用户填写并提交问卷', '提交成功'],
        ['T08-04', 'B端管理员关闭问卷', '状态变为"已关闭"，不再接受填写'],
        ['T08-05', 'B端管理员查看问卷统计结果', '显示各题目选项分布'],
    ]),
    ('B9 办事服务管理', [
        ['T09-01', 'C端用户查看办事事项分类列表', '按分类展示（教育/住房/医疗/就业/社保/交通/税务/证件）'],
        ['T09-02', 'C端用户查看事项详情', '显示事项编码、名称、描述、办理部门等'],
        ['T09-03', 'B端管理员新增办事事项', '新增成功'],
        ['T09-04', 'B端管理员编辑办事事项', '编辑保存成功'],
        ['T09-05', 'B端管理员删除办事事项', '删除成功（逻辑删除）'],
        ['T09-06', 'B端管理员按部门筛选事项', '仅显示对应部门的事项'],
    ]),
    ('B10 AI智能问答与搜索', [
        ['T10-01', 'C端用户发送问题，AI流式返回回答', 'SSE流式输出，回答内容正确'],
        ['T10-02', 'C端用户查看会话历史', '返回历史对话记录'],
        ['T10-03', 'C端用户通过关键词搜索', '返回匹配结果列表'],
        ['T10-04', 'C端用户使用语义搜索', '返回语义匹配结果'],
        ['T10-05', 'C端用户查看热门搜索关键词', '返回热门关键词列表'],
        ['T10-06', 'B端管理员查看AI审计列表', '分页展示AI问答审计记录'],
    ]),
    ('B11 绩效统计', [
        ['T11-01', 'B端管理员查看仪表盘数据', '显示访问量、服务量、留言量、满意度等核心指标'],
        ['T11-02', 'B端管理员查看周趋势图', '显示周一至周日访问量和服务量趋势'],
        ['T11-03', 'B端管理员按日期范围查询统计报表', '返回对应日期范围的统计数据'],
    ]),
]

for section_name, cases in test_cases:
    add_heading_styled(doc, section_name, 2)
    make_table(doc, ['编号', '测试要点', '预期结果'], cases, col_widths=[2, 7, 6])
    doc.add_paragraph()

# ========== 保存 ==========
output_path = r'e:\huadee\工程文档-指导版\6.测试计划.docx'
doc.save(output_path)
print(f'文档已保存至: {output_path}')
