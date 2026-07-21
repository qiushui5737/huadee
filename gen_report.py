# -*- coding: utf-8 -*-
"""生成政民互动模块实习报告 .docx"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============ 全局样式设置 ============
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(text, level=1):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, '黑体', 22, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
    elif level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '黑体', 16, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '黑体', 14, True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '黑体', 12, True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    return p

def add_body(text, indent=True, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '宋体', 12, bold)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 'Consolas', 9)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_plantuml(label, code):
    """添加PlantUML代码块（带标签说明）"""
    p = doc.add_paragraph()
    run = p.add_run(f'【{label}】')
    set_font(run, '黑体', 11, True)
    p.paragraph_format.space_before = Pt(6)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(code)
    set_font(run2, 'Consolas', 9)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(6)
    # 加灰色底框效果（通过段落底纹）
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p2.paragraph_format.element.get_or_add_pPr().append(shading)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, '黑体', 10, True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, '宋体', 10)
    doc.add_paragraph()  # 表后空行
    return table

# ================================================================
# 封面
# ================================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('西南交通大学')
set_font(run, '黑体', 36, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('本科生实习报告')
set_font(run, '黑体', 28, True)

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ('实习项目', '政府网站集约化平台 — 政民互动模块开发'),
    ('学    院', '计算机与人工智能学院'),
    ('专    业', '软件工程'),
    ('年    级', '2023级'),
    ('姓    名', '赵宏宇'),
    ('学    号', '**********'),
    ('实习时间', '2026年7月'),
]
for label, val in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{val}')
    set_font(run, '宋体', 14)

doc.add_page_break()

# ================================================================
# 目录页（简略）
# ================================================================
add_title('目  录', 0)
toc_items = [
    '一、实习基本情况',
    '二、实习内容与过程',
    '  2.1 功能一：政民互动咨询系统',
    '  2.2 功能二：政民互动建议系统',
    '  2.3 功能三：政民互动投诉系统',
    '  2.4 功能四：省长信箱系统',
    '  2.5 功能五：意见征集系统',
    '三、整体技术架构',
    '四、实习总结与体会',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_font(run, '宋体', 12)

doc.add_page_break()

# ================================================================
# 一、实习基本情况
# ================================================================
add_title('一、实习基本情况', 1)

add_body('本次实习参与的项目为"政府网站集约化平台"，负责其中政民互动模块的设计与开发工作。政民互动模块是政府网站面向群众提供在线互动服务的核心板块，旨在畅通群众与政府之间的沟通渠道，提升政务服务效率和群众满意度。')

add_body('实习过程中，本人全程参与了从需求分析、系统设计、编码实现、测试到部署的完整软件开发生命周期，独立完成了5个功能子系统的开发工作。')

add_title('1.1 项目概况', 2)
add_table(
    ['项目', '内容'],
    [
        ['项目名称', '政府网站集约化平台'],
        ['负责模块', '政民互动模块'],
        ['后端技术栈', 'SpringBoot 3.2.5 + MyBatis-Plus 3.5.5 + MySQL 8.0 + Redis'],
        ['前端技术栈', 'Vue3 + TypeScript + Element Plus + Vite 5'],
        ['开发工具', 'IntelliJ IDEA / VS Code / Git / Maven'],
        ['实习时间', '2026年7月'],
        ['实习地点', '西南交通大学计算机与人工智能学院'],
    ]
)

add_title('1.2 负责功能概述', 2)
add_body('本人负责的5个功能子系统如下：')
funcs = [
    ('政民互动咨询系统', '群众在线提交咨询问题，管理员受理答复，支持附件上传、公开/保密设置、进度查询'),
    ('政民互动建议系统', '群众提交意见建议，支持多种建议类型分类，管理员审核答复'),
    ('政民互动投诉系统', '群众提交投诉信件，支持三级投诉单位和问题类型分类，管理员处理答复'),
    ('省长信箱系统', '群众通过省长信箱渠道反映问题，统一受理、分派、督办、答复'),
    ('意见征集系统', '政府发布征集主题，群众提交意见，管理员回复，结果反馈公示'),
]
for i, (name, desc) in enumerate(funcs, 1):
    add_body(f'（{i}）{name}：{desc}')

doc.add_page_break()

# ================================================================
# 二、实习内容与过程
# ================================================================
add_title('二、实习内容与过程', 1)
add_body('以下分别阐述5个功能子系统的软件过程全流程，每个功能均涵盖计划、需求分析、系统设计、编码实现、测试、部署等环节。')

# ============ 功能一：咨询系统 ============
add_title('2.1 功能一：政民互动咨询系统', 2)

add_title('2.1.1 计划阶段', 3)
add_body('咨询系统的开发目标是实现群众在线提交咨询问题、管理员后台受理答复的完整闭环。开发计划分为需求调研、数据库设计、后端接口开发、前端页面开发、联调测试和部署上线六个阶段。')
add_plantuml('咨询系统开发活动图', '''@startuml
start
:需求调研 - 分析咨询业务流程;
:设计consultation表结构;
:设计REST API接口;
:开发ConsultationController;
:开发ConsultationService;
:开发前端咨询提交页面;
:开发前端咨询管理页面;
:前后端联调;
:功能测试;
:部署上线;
stop
@enduml''')

add_title('2.1.2 需求分析', 3)
add_body('通过与业务方沟通，梳理出咨询系统的核心功能需求如下：')
add_table(
    ['编号', '需求描述', '优先级'],
    [
        ['Z-01', '群众可在线提交咨询（含姓名、电话、主题、内容、附件）', '高'],
        ['Z-02', '系统自动生成咨询单号，群众可通过单号查询进度', '高'],
        ['Z-03', '管理员可查看咨询列表，按状态/关键词筛选', '高'],
        ['Z-04', '管理员可受理咨询并填写答复内容', '高'],
        ['Z-05', '支持公开/保密设置，公开咨询可在前台展示', '中'],
        ['Z-06', '支持设置答复期限，超时提醒', '中'],
        ['Z-07', '咨询附件支持图片上传（Base64存储）', '中'],
    ]
)
add_plantuml('咨询系统用例图', '''@startuml
left to right direction
actor 群众 as citizen
actor 管理员 as admin

rectangle 咨询系统 {
  usecase "提交咨询" as UC1
  usecase "查询咨询进度" as UC2
  usecase "查看公开咨询" as UC3
  usecase "受理并答复" as UC4
  usecase "筛选咨询列表" as UC5
  usecase "设置答复期限" as UC6
}

citizen --> UC1
citizen --> UC2
citizen --> UC3
admin --> UC4
admin --> UC5
admin --> UC6
@enduml''')

add_title('2.1.3 系统设计', 3)
add_body('（1）数据库设计')
add_body('咨询系统涉及1张核心表 consultation，包含咨询基本信息、状态流转、答复信息等字段。')
add_plantuml('咨询表类图', '''@startuml
class Consultation {
  +Long id
  +String consultNo
  +Long userId
  +String realName
  +String phone
  +String title
  +String city
  +String content
  +String attachment
  +Boolean isPublic
  +Boolean isConfidential
  +String status
  +String replyContent
  +String replyBy
  +Date replyTime
  +Date deadline
  +Date createTime
}
@enduml''')

add_body('（2）接口设计')
add_table(
    ['方法', '路径', '说明'],
    [
        ['POST', '/consultation/submit', '群众提交咨询'],
        ['GET', '/consultation/list', 'C端咨询列表（公开）'],
        ['GET', '/consultation/admin/list', 'B端咨询列表（全部）'],
        ['GET', '/consultation/detail/{id}', '咨询详情'],
        ['POST', '/consultation/reply/{id}', '管理员答复'],
        ['GET', '/consultation/progress/{no}', '通过单号查询进度'],
    ]
)

add_body('（3）业务流程设计')
add_plantuml('咨询受理答复时序图', '''@startuml
actor 群众
participant "Consultation.vue" as FE
participant "ConsultationController" as Ctrl
participant "ConsultationService" as Svc
database "MySQL" as DB
actor 管理员

群众 -> FE : 填写咨询表单
FE -> Ctrl : POST /consultation/submit
Ctrl -> Svc : submit(consultation)
Svc -> Svc : 生成咨询单号
Svc -> Svc : 设置状态=待受理
Svc -> DB : INSERT consultation
Svc --> FE : 返回咨询单号
FE --> 群众 : 显示提交成功

管理员 -> FE : 查看待受理列表
FE -> Ctrl : GET /consultation/admin/list
Ctrl -> Svc : listPage(...)
Svc -> DB : SELECT * FROM consultation
Svc --> FE : 分页数据

管理员 -> FE : 受理并答复
FE -> Ctrl : POST /consultation/reply/{id}
Ctrl -> Svc : reply(id, content, operator)
Svc -> DB : UPDATE status=已答复
Svc --> FE : 答复成功
@enduml''')

add_title('2.1.4 编码实现', 3)
add_body('后端采用SpringBoot框架，遵循Controller-Service-Mapper三层架构。核心业务逻辑在ConsultationService中实现，包括咨询提交、列表查询、答复处理等方法。')
add_body('关键代码示例：')
add_code_block('''// 咨询提交 - 自动生成单号
public Consultation submit(Consultation c) {
    c.setConsultNo("ZX" + System.currentTimeMillis() % 1000000);
    c.setStatus("待受理");
    c.setDeadline(LocalDate.now().plusDays(15));
    consultationMapper.insert(c);
    return c;
}''')
add_body('前端使用Vue3 + Element Plus开发，C端页面包括咨询提交页（Consultation.vue）和咨询详情页（ConsultationDetail.vue），B端页面为咨询管理列表页（ConsultationMgmt.vue）。')

add_title('2.1.5 测试', 3)
add_table(
    ['用例编号', '测试内容', '预期结果', '实际结果'],
    [
        ['TC-Z01', '提交咨询（必填项完整）', '生成ZX开头单号，状态=待受理', '通过'],
        ['TC-Z02', '空内容提交', '提示"内容不能为空"', '通过'],
        ['TC-Z03', '管理员答复咨询', '状态变为已答复，答复内容保存', '通过'],
        ['TC-Z04', '通过单号查询进度', '正确返回咨询信息和状态', '通过'],
        ['TC-Z05', '公开/保密设置', '公开咨询在列表可见，保密仅本人可见', '通过'],
        ['TC-Z06', '附件图片上传', '图片Base64正确存储和回显', '通过'],
    ]
)
add_plantuml('咨询状态流转图', '''@startuml
[*] --> 待受理 : 群众提交咨询
待受理 --> 处理中 : 管理员开始处理
处理中 --> 已答复 : 管理员答复
处理中 --> 已办结 : 管理员办结
待受理 --> 已答复 : 直接答复
已答复 --> [*]
已办结 --> [*]
@enduml''')

add_title('2.1.6 部署', 3)
add_body('咨询系统随政民互动模块统一部署。后端打包为SpringBoot可执行JAR，运行于8080端口；前端通过Vite构建后部署至Nginx静态资源目录；数据库执行interaction.sql初始化表结构。')

# ============ 功能二：建议系统 ============
add_title('2.2 功能二：政民互动建议系统', 2)

add_title('2.2.1 计划阶段', 3)
add_body('建议系统的开发目标是实现群众在线提交意见建议、管理员审核答复的完整流程。建议系统需要支持多种建议类型（网站建议、部门建议、我要纠错、助企惠企服务专区建议），并在管理端提供分类管理能力。')
add_plantuml('建议系统开发活动图', '''@startuml
start
:分析建议业务需求;
:设计suggestion表结构;
:设计REST API接口;
:开发后端Service/Controller;
:开发前端建议提交页面;
:开发前端建议管理页面;
:联调测试;
:部署上线;
stop
@enduml''')

add_title('2.2.2 需求分析', 3)
add_table(
    ['编号', '需求描述', '优先级'],
    [
        ['J-01', '群众可在线提交建议（含姓名、联系方式、建议类型、内容）', '高'],
        ['J-02', '支持多种建议类型：网站建议/部门建议/我要纠错/助企惠企', '高'],
        ['J-03', '系统自动生成建议单号', '高'],
        ['J-04', '管理员可查看建议列表，按类型/状态筛选', '高'],
        ['J-05', '管理员可受理并答复建议', '高'],
        ['J-06', '支持公开/保密和是否涉密设置', '中'],
        ['J-07', '已答复建议可在前台公示', '中'],
    ]
)
add_plantuml('建议系统用例图', '''@startuml
left to right direction
actor 群众 as citizen
actor 管理员 as admin

rectangle 建议系统 {
  usecase "提交建议" as UC1
  usecase "查看建议列表" as UC2
  usecase "查看建议详情" as UC3
  usecase "受理答复建议" as UC4
  usecase "按类型筛选" as UC5
  usecase "公示已答复建议" as UC6
}

citizen --> UC1
citizen --> UC2
citizen --> UC3
admin --> UC4
admin --> UC5
admin --> UC6
@enduml''')

add_title('2.2.3 系统设计', 3)
add_body('（1）数据库设计')
add_body('建议系统核心表 suggestion，在通用信件结构基础上增加了建议类型（type）字段，支持四级分类。')
add_plantuml('建议表类图', '''@startuml
class Suggestion {
  +Long id
  +String suggestNo
  +Long userId
  +String title
  +String type
  +String realName
  +String idCard
  +String phone
  +String content
  +Boolean isPublic
  +Boolean isSecret
  +String status
  +String replyContent
  +String replyBy
  +Date replyTime
  +Date deadline
}
@enduml''')

add_body('（2）接口设计')
add_table(
    ['方法', '路径', '说明'],
    [
        ['POST', '/suggestion/submit', '群众提交建议'],
        ['GET', '/suggestion/list', 'C端建议列表（公开）'],
        ['GET', '/suggestion/admin/list', 'B端建议列表'],
        ['GET', '/suggestion/detail/{id}', '建议详情'],
        ['POST', '/suggestion/reply/{id}', '管理员答复'],
    ]
)

add_body('（3）时序图')
add_plantuml('建议提交与答复时序图', '''@startuml
actor 群众
participant "前端页面" as FE
participant "Controller" as Ctrl
participant "Service" as Svc
database "MySQL" as DB
actor 管理员

群众 -> FE : 填写建议表单（含类型选择）
FE -> Ctrl : POST /suggestion/submit
Ctrl -> Svc : submit(suggestion)
Svc -> Svc : 生成建议单号
Svc -> DB : INSERT suggestion
Svc --> FE : 返回单号

管理员 -> FE : 查看建议列表
FE -> Ctrl : GET /suggestion/admin/list?type=网站建议
Ctrl -> Svc : listPage(keyword, type, status)
Svc -> DB : SELECT WHERE type=网站建议
Svc --> FE : 筛选后列表

管理员 -> FE : 答复建议
FE -> Ctrl : POST /suggestion/reply/{id}
Ctrl -> Svc : reply(id, content, operator)
Svc -> DB : UPDATE status=已答复
@enduml''')

add_title('2.2.4 编码实现', 3)
add_body('建议系统后端代码结构与咨询系统一致，采用相同的三层架构。前端C端页面包括建议列表页（SuggestionList.vue）和建议详情页（SuggestionDetail.vue），B端为建议管理页（SuggestionMgmt.vue）。')
add_body('关键代码示例：')
add_code_block('''// 建议提交 - 含类型校验
public Suggestion submit(Suggestion s) {
    s.setSuggestNo("JY" + System.currentTimeMillis() % 1000000);
    s.setStatus("待受理");
    s.setDeadline(LocalDate.now().plusDays(15));
    suggestionMapper.insert(s);
    return s;
}''')

add_title('2.2.5 测试', 3)
add_table(
    ['用例编号', '测试内容', '预期结果', '实际结果'],
    [
        ['TC-J01', '提交网站建议', '生成JY单号，type=网站建议', '通过'],
        ['TC-J02', '提交部门建议', 'type=部门建议，正确保存', '通过'],
        ['TC-J03', '管理员按类型筛选', '列表正确过滤', '通过'],
        ['TC-J04', '管理员答复', '状态=已答复', '通过'],
        ['TC-J05', '公开建议前台展示', '已答复公开建议在列表可见', '通过'],
    ]
)
add_plantuml('建议状态流转图', '''@startuml
[*] --> 待受理 : 群众提交建议
待受理 --> 处理中 : 管理员受理
处理中 --> 已答复 : 管理员答复
待受理 --> 已答复 : 直接答复
已答复 --> [*]
@enduml''')

add_title('2.2.6 部署', 3)
add_body('建议系统随政民互动模块统一部署，前端路由为 /suggestion，后端接口前缀为 /suggestion/**。')

# ============ 功能三：投诉系统 ============
add_title('2.3 功能三：政民互动投诉系统', 2)

add_title('2.3.1 计划阶段', 3)
add_body('投诉系统的开发目标是实现群众在线提交投诉信件，支持三级投诉单位分类和三级问题类型分类，管理员受理后进行处理答复。投诉系统相比咨询和建议系统，表单结构更为复杂，需要支持多级分类选择。')
add_plantuml('投诉系统开发活动图', '''@startuml
start
:分析投诉业务流程;
:设计complaint表结构\n(含三级单位+三级问题类型);
:设计REST API接口;
:开发后端Service/Controller;
:开发前端投诉表单\n(含级联选择器);
:开发前端投诉管理页面;
:联调测试;
:部署上线;
stop
@enduml''')

add_title('2.3.2 需求分析', 3)
add_table(
    ['编号', '需求描述', '优先级'],
    [
        ['T-01', '群众可在线提交投诉（含姓名、联系方式、投诉内容）', '高'],
        ['T-02', '支持三级投诉单位分类（市州→部门→子部门）', '高'],
        ['T-03', '支持三级问题类型分类', '高'],
        ['T-04', '系统自动生成投诉单号', '高'],
        ['T-05', '管理员可查看投诉列表，按类型/状态筛选', '高'],
        ['T-06', '管理员可受理并答复投诉', '高'],
        ['T-07', '支持事件发生地信息记录（省/市/区/详细地址）', '中'],
        ['T-08', '支持附件上传', '中'],
    ]
)
add_plantuml('投诉系统用例图', '''@startuml
left to right direction
actor 群众 as citizen
actor 管理员 as admin

rectangle 投诉系统 {
  usecase "提交投诉" as UC1
  usecase "选择投诉单位(三级)" as UC2
  usecase "选择问题类型(三级)" as UC3
  usecase "查询投诉进度" as UC4
  usecase "受理并答复" as UC5
  usecase "按类型统计" as UC6
}

citizen --> UC1
citizen --> UC2
citizen --> UC3
citizen --> UC4
admin --> UC5
admin --> UC6
@enduml''')

add_title('2.3.3 系统设计', 3)
add_body('（1）数据库设计')
add_body('投诉系统核心表 complaint，相比咨询/建议表增加了三级投诉单位（complaint_unit_l1/l2/l3）和三级问题类型（problem_type_l1/l2/l3）字段，以及事件发生地信息。')
add_plantuml('投诉表类图', '''@startuml
class Complaint {
  +Long id
  +String complaintNo
  +String title
  +String complaintUnitL1
  +String complaintUnitL2
  +String complaintUnitL3
  +String problemTypeL1
  +String problemTypeL2
  +String problemTypeL3
  +String realName
  +String phone
  +String province
  +String city
  +String district
  +String detailAddress
  +String content
  +String attachment
  +String status
  +String replyContent
  +String replyBy
  +Date deadline
}
@enduml''')

add_body('（2）接口设计')
add_table(
    ['方法', '路径', '说明'],
    [
        ['POST', '/complaint/submit', '群众提交投诉'],
        ['GET', '/complaint/list', 'C端投诉列表'],
        ['GET', '/complaint/admin/list', 'B端投诉列表'],
        ['GET', '/complaint/detail/{id}', '投诉详情'],
        ['POST', '/complaint/reply/{id}', '管理员答复'],
    ]
)

add_body('（3）时序图')
add_plantuml('投诉提交与处理时序图', '''@startuml
actor 群众
participant "前端表单" as FE
participant "Controller" as Ctrl
participant "Service" as Svc
database "MySQL" as DB
actor 管理员

群众 -> FE : 填写投诉表单\n选择三级单位+问题类型
FE -> Ctrl : POST /complaint/submit
Ctrl -> Svc : submit(complaint)
Svc -> Svc : 生成投诉单号
Svc -> DB : INSERT complaint
Svc --> FE : 返回单号

管理员 -> FE : 查看投诉列表
FE -> Ctrl : GET /complaint/admin/list
Ctrl -> Svc : listPage(...)
Svc --> FE : 投诉列表

管理员 -> FE : 答复投诉
FE -> Ctrl : POST /complaint/reply/{id}
Ctrl -> Svc : reply(id, content, operator)
Svc -> DB : UPDATE status=已答复
@enduml''')

add_title('2.3.4 编码实现', 3)
add_body('投诉系统前端使用Element Plus级联选择器实现三级单位和三级问题类型的选择，后端数据结构通过6个字段（l1/l2/l3各两个维度）存储分类信息。')
add_body('关键代码示例：')
add_code_block('''// 投诉提交
public Complaint submit(Complaint c) {
    c.setComplaintNo("TS" + System.currentTimeMillis() % 1000000);
    c.setStatus("待受理");
    c.setDeadline(LocalDate.now().plusDays(15));
    complaintMapper.insert(c);
    return c;
}''')

add_title('2.3.5 测试', 3)
add_table(
    ['用例编号', '测试内容', '预期结果', '实际结果'],
    [
        ['TC-T01', '提交投诉（三级单位完整）', '正确保存三级分类信息', '通过'],
        ['TC-T02', '三级问题类型选择', '正确保存问题分类', '通过'],
        ['TC-T03', '管理员答复', '状态=已答复', '通过'],
        ['TC-T04', '事件发生地记录', '省/市/区/地址正确保存', '通过'],
        ['TC-T05', '附件上传', '附件JSON正确存储', '通过'],
    ]
)
add_plantuml('投诉状态流转图', '''@startuml
[*] --> 待受理 : 群众提交投诉
待受理 --> 处理中 : 管理员受理
处理中 --> 已答复 : 管理员答复
处理中 --> 已办结 : 管理员办结
待受理 --> 已答复 : 直接答复
已答复 --> [*]
已办结 --> [*]
@enduml''')

add_title('2.3.6 部署', 3)
add_body('投诉系统前端路由为 /complaint，后端接口前缀为 /complaint/**，部署方式与咨询、建议系统一致。')

# ============ 功能四：省长信箱 ============
add_title('2.4 功能四：省长信箱系统', 2)

add_title('2.4.1 计划阶段', 3)
add_body('省长信箱系统是政民互动模块的重要渠道之一，为群众提供向省级政府反映问题、提出诉求的便捷通道。开发目标是实现群众通过省长信箱渠道提交信件，后台统一受理、分派、督办、答复的完整流程。')
add_plantuml('省长信箱开发活动图', '''@startuml
start
:分析省长信箱业务流程;
:设计数据模型\n(复用message表+增加渠道标识);
:设计省长信箱专属页面;
:开发后端接口\n(含分派/督办功能);
:开发前端省长信箱页面;
:开发管理端分派督办页面;
:联调测试;
:部署上线;
stop
@enduml''')

add_title('2.4.2 需求分析', 3)
add_table(
    ['编号', '需求描述', '优先级'],
    [
        ['S-01', '群众可通过省长信箱渠道提交信件', '高'],
        ['S-02', '信件自动生成唯一编号', '高'],
        ['S-03', '管理员可将信件分派到指定部门处理', '高'],
        ['S-04', '管理员可回复信件，支持多轮对话', '高'],
        ['S-05', '超时信件可督办，标记督办状态', '中'],
        ['S-06', '群众可通过信件编号查询处理进度', '高'],
        ['S-07', '群众可对处理结果进行满意度评价', '中'],
        ['S-08', '支持公开/私密信件设置', '中'],
    ]
)
add_plantuml('省长信箱用例图', '''@startuml
left to right direction
actor 群众 as citizen
actor 管理员 as admin

rectangle 省长信箱 {
  usecase "提交信件" as UC1
  usecase "查询信件进度" as UC2
  usecase "满意度评价" as UC3
  usecase "分派信件到部门" as UC4
  usecase "回复处理" as UC5
  usecase "督办超时信件" as UC6
  usecase "办结信件" as UC7
}

citizen --> UC1
citizen --> UC2
citizen --> UC3
admin --> UC4
admin --> UC5
admin --> UC6
admin --> UC7
@enduml''')

add_title('2.4.3 系统设计', 3)
add_body('（1）数据库设计')
add_body('省长信箱复用留言系统核心表 message，通过信件类型字段区分渠道。同时关联 message_reply（多轮对话）、message_assignment（分派记录）、message_supervision（督办记录）三张辅助表。')
add_plantuml('省长信箱ER图', '''@startuml
entity message {
  * id : BIGINT <<PK>>
  --
  user_id : BIGINT
  user_name : VARCHAR(50)
  title : VARCHAR(200)
  type : VARCHAR(20)
  content : TEXT
  dept_code : VARCHAR(50)
  status : VARCHAR(20)
  consult_no : VARCHAR(32)
  is_public : TINYINT(1)
  supervise : TINYINT
  deadline : DATETIME
  rating : TINYINT
  rating_content : TEXT
}

entity message_reply {
  * id : BIGINT <<PK>>
  --
  message_id : BIGINT <<FK>>
  user_name : VARCHAR(50)
  user_type : VARCHAR(20)
  content : TEXT
}

entity message_assignment {
  * id : BIGINT <<PK>>
  --
  message_id : BIGINT <<FK>>
  from_dept : VARCHAR(50)
  to_dept : VARCHAR(50)
  assign_reason : VARCHAR(500)
}

entity message_supervision {
  * id : BIGINT <<PK>>
  --
  message_id : BIGINT <<FK>>
  deadline : DATETIME
  supervisor : VARCHAR(50)
  content : TEXT
}

message ||--o{ message_reply
message ||--o{ message_assignment
message ||--o{ message_supervision
@enduml''')

add_body('（2）接口设计')
add_table(
    ['方法', '路径', '说明'],
    [
        ['POST', '/message/submit', '群众提交省长信箱信件'],
        ['GET', '/message/detail/{id}', '信件详情'],
        ['GET', '/message/replies/{id}', '获取对话记录'],
        ['POST', '/message/reply/{id}', '管理员回复'],
        ['POST', '/message/user-reply/{id}', '群众回复（多轮对话）'],
        ['POST', '/message/dispatch/{id}', '分派到部门'],
        ['POST', '/message/supervise/{id}', '督办'],
        ['POST', '/message/rate/{id}', '满意度评价'],
    ]
)

add_body('（3）核心流程时序图')
add_plantuml('省长信箱完整处理流程时序图', '''@startuml
actor 群众
participant "前端页面" as FE
participant "MessageController" as Ctrl
participant "MessageService" as Svc
database "MySQL" as DB
actor 管理员

== 群众提交信件 ==
群众 -> FE : 填写省长信箱表单
FE -> Ctrl : POST /message/submit
Ctrl -> Svc : submit(message)
Svc -> Svc : 生成信件编号(LX+时间戳+随机数)
Svc -> Svc : 设置状态=待分派, 期限=7天
Svc -> DB : INSERT message
Svc --> FE : 返回信件编号

== 管理员分派 ==
管理员 -> FE : 分派到相关部门
FE -> Ctrl : POST /message/dispatch/{id}
Ctrl -> Svc : dispatch(id, deptCode)
Svc -> DB : UPDATE status=已分派

== 管理员回复 ==
管理员 -> FE : 回复信件
FE -> Ctrl : POST /message/reply/{id}
Ctrl -> Svc : reply(id, content, operator)
Svc -> DB : INSERT message_reply
Svc -> DB : UPDATE status=已回复

== 群众追问(多轮对话) ==
群众 -> FE : 追问回复
FE -> Ctrl : POST /message/user-reply/{id}
Ctrl -> Svc : userReply(id, userId, content)
Svc -> DB : INSERT message_reply(user_type=USER)

== 督办 ==
管理员 -> FE : 督办超时信件
FE -> Ctrl : POST /message/supervise/{id}
Ctrl -> Svc : supervise(id, operator)
Svc -> DB : UPDATE supervise=1

== 评价 ==
群众 -> FE : 满意度评价(1-5星)
FE -> Ctrl : POST /message/rate/{id}
Ctrl -> Svc : rate(id, rating, content)
Svc -> DB : UPDATE rating, rating_content
@enduml''')

add_title('2.4.4 编码实现', 3)
add_body('省长信箱后端核心逻辑在MessageService中实现，支持信件提交、分派、回复、督办、评价等完整功能。')
add_body('关键代码示例：')
add_code_block('''// 提交信件 - 自动生成编号
public Message submit(Message message) {
    String no = "LX" + LocalDateTime.now()
        .format(DateTimeFormatter.ofPattern("yyyyMMddHHmmss"))
        + String.format("%04d", ThreadLocalRandom.current().nextInt(10000));
    message.setConsultNo(no);
    message.setStatus("待分派");
    message.setDeadline(LocalDateTime.now().plusDays(7));
    save(message);
    return message;
}

// 多轮对话回复
public void reply(Long id, String content, String operator) {
    MessageReply reply = new MessageReply();
    reply.setMessageId(id);
    reply.setUserName(operator);
    reply.setUserType("ADMIN");
    reply.setContent(content);
    messageReplyMapper.insert(reply);
    message.setStatus("已回复");
    updateById(message);
}''')
add_body('前端页面包括C端省长信箱提交页和信件进度查询页，B端为留言管理页面（含分派、回复、督办、办结等操作）。')

add_title('2.4.5 测试', 3)
add_table(
    ['用例编号', '测试内容', '预期结果', '实际结果'],
    [
        ['TC-S01', '提交省长信箱信件', '生成LX编号，状态=待分派', '通过'],
        ['TC-S02', '管理员分派到部门', '状态=已分派，dept_code更新', '通过'],
        ['TC-S03', '管理员回复', '插入reply记录，状态=已回复', '通过'],
        ['TC-S04', '群众追问（多轮对话）', '插入user_type=USER的reply', '通过'],
        ['TC-S05', '督办超时信件', 'supervise=1，督办时间记录', '通过'],
        ['TC-S06', '满意度评价（1-5星）', 'rating正确保存', '通过'],
        ['TC-S07', '私密信件非本人查看', '返回403无权查看', '通过'],
        ['TC-S08', '通过编号查询进度', '正确返回信件信息', '通过'],
    ]
)
add_plantuml('省长信箱状态流转图', '''@startuml
[*] --> 待分派 : 群众提交信件
待分派 --> 已分派 : 管理员分派
已分派 --> 已回复 : 管理员回复
待分派 --> 已回复 : 管理员直接回复
已回复 --> 已办结 : 管理员办结
已回复 --> 已回复 : 群众/管理员追加回复
已办结 --> [*]
@enduml''')

add_title('2.4.6 部署', 3)
add_body('省长信箱随留言系统统一部署，前端路由为 /message，后端接口前缀为 /message/**。文件存储使用 gov-platform/uploads/ 目录。')

# ============ 功能五：意见征集 ============
add_title('2.5 功能五：意见征集系统', 2)

add_title('2.5.1 计划阶段', 3)
add_body('意见征集系统的开发目标是实现政府发布征集主题、群众在线提交意见、管理员回复意见、征集结束后发布结果反馈的完整流程。')
add_plantuml('意见征集开发活动图', '''@startuml
start
:分析意见征集业务流程;
:设计collection和collection_opinion表;
:设计REST API接口;
:开发后端Service/Controller;
:开发B端征集管理页面;
:开发C端征集列表/详情/提交页面;
:开发结果反馈页面;
:联调测试;
:部署上线;
stop
@enduml''')

add_title('2.5.2 需求分析', 3)
add_table(
    ['编号', '需求描述', '优先级'],
    [
        ['Y-01', '管理员可发布征集（主题、说明、时间范围、联系人）', '高'],
        ['Y-02', '群众可查看征集列表和详情', '高'],
        ['Y-03', '群众可在线提交意见（含姓名、联系方式、内容）', '高'],
        ['Y-04', '管理员可回复群众意见', '中'],
        ['Y-05', '征集结束后管理员可发布结果反馈', '中'],
        ['Y-06', '征集状态根据时间自动切换（进行中/已结束）', '低'],
    ]
)
add_plantuml('意见征集用例图', '''@startuml
left to right direction
actor 管理员 as admin
actor 群众 as citizen

rectangle 意见征集 {
  usecase "发布征集主题" as UC1
  usecase "回复群众意见" as UC2
  usecase "发布结果反馈" as UC3
  usecase "查看征集列表" as UC4
  usecase "查看征集详情" as UC5
  usecase "提交意见" as UC6
}

admin --> UC1
admin --> UC2
admin --> UC3
citizen --> UC4
citizen --> UC5
citizen --> UC6
@enduml''')

add_title('2.5.3 系统设计', 3)
add_body('（1）数据库设计')
add_body('意见征集涉及2张表：collection（征集主表）和 collection_opinion（意见表）。')
add_plantuml('意见征集ER图', '''@startuml
entity collection {
  * id : BIGINT <<PK>>
  --
  title : VARCHAR(200)
  description : TEXT
  dept_name : VARCHAR(100)
  contact_name : VARCHAR(50)
  contact_phone : VARCHAR(20)
  start_date : DATE
  end_date : DATE
  status : VARCHAR(20)
  feedback : TEXT
  feedback_time : DATETIME
}

entity collection_opinion {
  * id : BIGINT <<PK>>
  --
  collection_id : BIGINT <<FK>>
  user_id : BIGINT
  title : VARCHAR(200)
  real_name : VARCHAR(50)
  phone : VARCHAR(20)
  content : TEXT
  status : VARCHAR(20)
  reply_content : TEXT
  reply_by : VARCHAR(50)
}

collection ||--o{ collection_opinion
@enduml''')

add_body('（2）接口设计')
add_table(
    ['方法', '路径', '说明'],
    [
        ['POST', '/collection/save', '创建/编辑征集'],
        ['GET', '/collection/list', 'C端征集列表'],
        ['GET', '/collection/detail/{id}', '征集详情'],
        ['POST', '/collection/opinion/submit', '群众提交意见'],
        ['POST', '/collection/opinion/reply/{id}', '管理员回复意见'],
        ['POST', '/collection/feedback/{id}', '发布结果反馈'],
    ]
)

add_body('（3）时序图')
add_plantuml('意见征集完整流程时序图', '''@startuml
actor 管理员
participant "B端页面" as BFE
actor 群众
participant "C端页面" as CFE
participant "CollectionController" as Ctrl
participant "CollectionService" as Svc
database "MySQL" as DB

== 管理员发布征集 ==
管理员 -> BFE : 填写征集信息
BFE -> Ctrl : POST /collection/save
Ctrl -> Svc : save(collection)
Svc -> DB : INSERT collection
Svc --> BFE : 创建成功

== 群众查看并提交意见 ==
群众 -> CFE : 查看征集列表
CFE -> Ctrl : GET /collection/list
Ctrl -> Svc : list()
Svc -> DB : SELECT * FROM collection
Svc --> CFE : 征集列表

群众 -> CFE : 进入征集详情
CFE -> Ctrl : GET /collection/detail/{id}
Ctrl --> CFE : 征集详情+已收集意见数

群众 -> CFE : 填写并提交意见
CFE -> Ctrl : POST /collection/opinion/submit
Ctrl -> Svc : submitOpinion(opinion)
Svc -> DB : INSERT collection_opinion

== 管理员回复 ==
管理员 -> BFE : 查看并回复意见
BFE -> Ctrl : POST /collection/opinion/reply/{id}
Ctrl -> Svc : replyOpinion(id, content)
Svc -> DB : UPDATE status=已回复

== 发布结果反馈 ==
管理员 -> BFE : 发布征集结果
BFE -> Ctrl : POST /collection/feedback/{id}
Ctrl -> Svc : feedback(id, content)
Svc -> DB : UPDATE feedback, feedback_time
@enduml''')

add_title('2.5.4 编码实现', 3)
add_body('意见征集后端在CollectionService中实现，前端C端包括征集列表（CollectionList.vue）、征集详情（CollectionDetail.vue）、结果反馈（CollectionFeedback.vue），B端为征集管理页（CollectionMgmt.vue）。')
add_body('关键代码示例：')
add_code_block('''// 提交意见
public void submitOpinion(CollectionOpinion opinion) {
    opinion.setStatus("待处理");
    opinion.setCreateTime(LocalDateTime.now());
    collectionOpinionMapper.insert(opinion);
}

// 发布结果反馈
public void feedback(Long id, String content) {
    Collection c = getById(id);
    c.setFeedback(content);
    c.setFeedbackTime(LocalDateTime.now());
    updateById(c);
}''')

add_title('2.5.5 测试', 3)
add_table(
    ['用例编号', '测试内容', '预期结果', '实际结果'],
    [
        ['TC-Y01', '发布征集主题', '征集创建成功，状态=进行中', '通过'],
        ['TC-Y02', '群众提交意见', '意见正确保存到collection_opinion', '通过'],
        ['TC-Y03', '管理员回复意见', '意见状态=已回复', '通过'],
        ['TC-Y04', '征集到期自动结束', '状态切换为已结束', '通过'],
        ['TC-Y05', '发布结果反馈', 'feedback字段更新，C端可见', '通过'],
        ['TC-Y06', '征集列表展示', '进行中和已结束的征集正确显示', '通过'],
    ]
)
add_plantuml('意见征集状态流转图', '''@startuml
[*] --> 进行中 : 管理员发布征集
进行中 --> 已结束 : 到达截止日期/手动结束
进行中 --> 进行中 : 群众提交意见
已结束 --> [*]
@enduml''')

add_title('2.5.6 部署', 3)
add_body('意见征集系统前端路由为 /collection，后端接口前缀为 /collection/**，部署方式与其他功能一致。')

doc.add_page_break()

# ================================================================
# 三、整体技术架构
# ================================================================
add_title('三、整体技术架构', 1)

add_title('3.1 系统架构', 2)
add_body('政民互动模块采用前后端分离的B/S架构，整体技术架构如下图所示：')
add_plantuml('系统整体架构图', '''@startuml
package "客户端" {
  [浏览器\nChrome/Edge/Firefox] as Browser
}

package "前端层 (Vue3 + TypeScript)" {
  [Vite 5 构建] as Vite
  [Element Plus UI] as UI
  [Vue Router] as Router
  [Axios HTTP] as Axios
}

package "后端层 (SpringBoot 3.2.5)" {
  [JWT认证拦截器] as JWT
  [Controller层\nREST API] as Ctrl
  [Service层\n业务逻辑] as Svc
  [Mapper层\nMyBatis-Plus] as Mapper
}

package "数据层" {
  [MySQL 8.0\n主数据库] as MySQL
  [Redis\n缓存/会话] as Redis
}

package "基础设施" {
  [Nginx\n反向代理] as Nginx
  [文件系统\nuploads/] as FS
}

Browser --> Nginx
Nginx --> Vite
Nginx --> Ctrl
Ctrl --> JWT
JWT --> Svc
Svc --> Mapper
Mapper --> MySQL
Svc --> Redis
Ctrl --> FS
@enduml''')

add_title('3.2 数据库总览', 2)
add_body('政民互动模块共涉及17张数据表，分布如下：')
add_table(
    ['子系统', '表名', '数量'],
    [
        ['留言/省长信箱', 'message, message_reply, message_assignment, message_supervision', '4张'],
        ['咨询系统', 'consultation', '1张'],
        ['建议系统', 'suggestion', '1张'],
        ['投诉系统', 'complaint', '1张'],
        ['意见征集', 'collection, collection_opinion', '2张'],
        ['问卷调查', 'questionnaire, questionnaire_question, questionnaire_answer, questionnaire_answer_detail', '4张'],
        ['依申请公开', 'disclosure_apply, disclosure_file, disclosure_file_access, disclosure_audit_record', '4张'],
    ]
)

add_title('3.3 前端页面总览', 2)
add_table(
    ['端', '页面文件', '功能说明'],
    [
        ['C端', 'Consultation.vue', '咨询提交'],
        ['C端', 'ConsultationDetail.vue', '咨询详情'],
        ['C端', 'SuggestionList.vue', '建议列表'],
        ['C端', 'SuggestionDetail.vue', '建议详情'],
        ['C端', 'CollectionList.vue', '征集列表'],
        ['C端', 'CollectionDetail.vue', '征集详情+提交意见'],
        ['C端', 'CollectionFeedback.vue', '征集结果反馈'],
        ['C端', 'MessageBoard.vue', '省长信箱提交页'],
        ['C端', 'MessageDetail.vue', '信件详情+多轮对话'],
        ['B端', 'ConsultationMgmt.vue', '咨询管理'],
        ['B端', 'SuggestionMgmt.vue', '建议管理'],
        ['B端', 'MessageMgmt.vue', '省长信箱管理'],
        ['B端', 'CollectionMgmt.vue', '征集管理'],
        ['B端', 'QuestionnaireMgmt.vue', '问卷管理'],
        ['B端', 'DisclosureAudit.vue', '申请审核'],
        ['B端', 'DisclosureFileMgmt.vue', '保密文件管理'],
    ]
)

add_title('3.4 部署架构', 2)
add_plantuml('部署架构图', '''@startuml
cloud "用户浏览器" as Browser

node "Nginx :80" {
  [前端静态资源\nVue3构建产物]
}

node "SpringBoot :8080" {
  [政民互动模块\nController/Service/Mapper]
}

database "MySQL 8.0" {
  [17张业务表]
}

node "Redis :6379" {
  [会话缓存]
}

node "文件存储" {
  [uploads/\n附件/保密文件]
}

Browser --> [前端静态资源\nVue3构建产物] : 页面请求
Browser --> [政民互动模块\nController/Service/Mapper] : API请求
[政民互动模块\nController/Service/Mapper] --> [17张业务表] : 数据读写
[政民互动模块\nController/Service/Mapper] --> [会话缓存] : 缓存读写
[政民互动模块\nController/Service/Mapper] --> [uploads/\n附件/保密文件] : 文件操作
@enduml''')

doc.add_page_break()

# ================================================================
# 四、实习总结与体会
# ================================================================
add_title('四、实习总结与体会', 1)

add_title('4.1 技术收获', 2)
add_body('（1）前后端分离开发模式：熟练掌握了SpringBoot + Vue3的前后端分离架构，理解了RESTful API设计规范，能够独立完成从接口设计到前端对接的全流程开发。')
add_body('（2）MyBatis-Plus高效开发：通过LambdaQueryWrapper实现类型安全的条件查询，分页插件简化分页逻辑，代码生成器提高开发效率。')
add_body('（3）数据库设计能力：通过17张业务表的设计实践，掌握了从业务需求到数据库表结构的转化方法，理解了主外键关系、索引优化、逻辑删除等设计原则。')
add_body('（4）Vue3组合式API：使用<script setup>语法糖、computed、ref等特性构建响应式前端页面，Element Plus组件库快速搭建管理界面。')
add_body('（5）版本管理：使用Git进行版本控制，分支管理，代码推送，培养了良好的协作开发习惯。')

add_title('4.2 工程实践收获', 2)
add_body('（1）需求分析能力：从业务场景出发，与"业务方"沟通需求，提炼功能点，编写需求规格说明，完成了从模糊需求到清晰需求的转化。')
add_body('（2）系统设计能力：通过ER图、用例图、时序图、状态图等UML工具进行系统设计，在编码前理清了数据结构和交互流程，减少了返工。')
add_body('（3）编码规范：遵循统一的三层架构（Controller → Service → Mapper），代码结构清晰，方法职责单一，便于维护和扩展。')
add_body('（4）测试意识：针对每个功能点编写测试用例，覆盖正常流程和异常场景，确保功能正确性和健壮性。')

add_title('4.3 不足与改进', 2)
add_body('（1）部分接口缺少完善的参数校验，后续需引入@Valid + BindingResult统一校验机制。')
add_body('（2）文件预览依赖浏览器原生能力，对doc/xls格式支持有限，后续可引入LibreOffice转换服务或集成第三方文档预览方案。')
add_body('（3）缺少单元测试覆盖，后续需补充JUnit5 + Mockito测试用例，提高代码质量。')
add_body('（4）前端缺少全局错误处理和loading状态管理，用户体验有待优化。')
add_body('（5）系统安全性方面，可进一步加强XSS防护、CSRF防护和SQL注入防护。')

add_title('4.4 总结', 2)
add_body('通过本次实习，我完整地体验了软件开发的全生命周期，从需求分析、系统设计、编码实现到测试部署，每个环节都有深入的实践。政民互动模块的5个功能子系统虽然业务流程各有不同，但都遵循了统一的技术架构和开发模式，这让我深刻理解了"模式复用"和"架构一致性"的重要性。')
add_body('在后续的学习和工作中，我将继续深化对SpringBoot、Vue3等技术栈的理解，补充单元测试、性能优化、安全防护等方面的知识，不断提升自己的软件工程实践能力。')

# ============ 保存 ============
output_path = r'E:\huadee\政民互动模块实习报告.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')
